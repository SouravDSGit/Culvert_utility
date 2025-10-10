"""
utils/subroutine_generate_CULVERT_Report.py
Create a new file called utils.py in your Flask application directory 
and add all the functions below:
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
from datetime import datetime
import ast

def generate_culvert_report(user_outputs_dir, project_name, user_name, current_date, current_time, 
                          ws_deln_response_file_path, hydro_vuln_response_file_path, 
                          hydro_geo_response_file_path, generated_plots):
    """
    Generate a comprehensive CULVERT report in DOCX format
    """
    
    try:
        # Create document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Read response files
        ws_deln_responses = read_response_file(ws_deln_response_file_path) if os.path.exists(ws_deln_response_file_path) else {}
        hydro_vuln_responses = read_response_file(hydro_vuln_response_file_path) if os.path.exists(hydro_vuln_response_file_path) else {}
        hydro_geo_responses = read_response_file(hydro_geo_response_file_path) if os.path.exists(hydro_geo_response_file_path) else {}
        
        # Add all sections
        add_title_page(doc, project_name, user_name, current_date, current_time)
        add_table_of_contents(doc)
        add_executive_summary(doc, project_name, user_outputs_dir, ws_deln_responses, hydro_vuln_responses, hydro_geo_responses)
        add_introduction(doc)
        add_study_area_overview(doc, generated_plots, project_name)
        add_data_summary_table(doc)
        add_results_section(doc, generated_plots, project_name)
        add_appendices(doc, ws_deln_responses, hydro_vuln_responses, hydro_geo_responses)
        
        # Save the document
        reports_dir = os.path.join(user_outputs_dir, 'Reports')
        os.makedirs(reports_dir, exist_ok=True)
        
        report_filename = f'CULVERT_{project_name}_report.docx'
        report_path = os.path.join(reports_dir, report_filename)
        doc.save(report_path)
        
        print(f"Report generated successfully: {report_path}")
        return report_path
        
    except Exception as e:
        print(f"Error generating DOCX report: {str(e)}")
        return None
    
def read_response_file(file_path):
    """Read and parse response file"""
    responses = {}
    try:
        with open(file_path, 'r') as f:
            for line in f:
                if ':' in line:
                    key, value = line.strip().split(':', 1)
                    if value.strip().startswith('['):
                        try:
                            responses[key.strip()] = ast.literal_eval(value.strip())
                        except:
                            responses[key.strip()] = value.strip()
                    else:
                        responses[key.strip()] = value.strip()
    except Exception as e:
        print(f"Error reading response file {file_path}: {str(e)}")
    return responses

def add_title_page(doc, project_name, user_name, current_date, current_time):
    """Add title page"""
    # Add title
    title = doc.add_heading('Climate and Upland Loading Vulnerability Evaluation and Risk Analysis Tool (CULVERT) Web Application: Version 1.0', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading(f'Hydrologic and Hydrogeomorphologic Risk Assessment Report', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add project info
    project_para = doc.add_paragraph()
    project_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    project_run = project_para.add_run(f'\nProject: {project_name}\n\n')
    project_run.font.size = Pt(16)
    project_run.bold = True
    
    # Add user and date info
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(f'Prepared by: {user_name}\n')
    info_run.font.size = Pt(12)
    
    date_run = info_para.add_run(f'Date: {current_date}\n')
    date_run.font.size = Pt(12)
    
    time_run = info_para.add_run(f'Time: {current_time}\n\n')
    time_run.font.size = Pt(12)
    
    # Add USDA Forest Service attribution
    attribution_para = doc.add_paragraph()
    attribution_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    attr_run = attribution_para.add_run('Generated using USDA Forest Service CULVERT Web Application- Version 1.0')
    attr_run.font.size = Pt(10)
    attr_run.italic = True
    
    # Add page break
    doc.add_page_break()

def add_table_of_contents(doc):
    """Add table of contents"""
    doc.add_heading('Table of Contents', 1)
    
    toc_items = [
    '1. Executive Summary',
    '2. Introduction', 
    '3. Study Area Overview',
    '4. Data Summary',
    '5. Results and Analysis',
    '5.1 Watershed and Road-Stream Crossing Analysis',
    '5.2 Hydrologic Risk Analysis',
    '5.3 Hydrogeomorphologic Risk Analysis',
    'Appendix A: Analysis Parameters',
    'Appendix B: Description of Methodologies',
    'Appendix C: References and Citations',
    'Appendix D: Glossary of Technical Terms'
    ]
    
    for item in toc_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()

def add_executive_summary(doc, project_name, user_outputs_dir, ws_deln_responses, hydro_vuln_responses, hydro_geo_responses):
    """Add dynamic executive summary based on analysis results"""
    import geopandas as gpd
    import os
    
    print("DEBUG: Starting add_executive_summary function")
    print(f"DEBUG: project_name = {project_name}")
    print(f"DEBUG: user_outputs_dir = {user_outputs_dir}")
    print(f"DEBUG: ws_deln_responses type = {type(ws_deln_responses)}")
    print(f"DEBUG: hydro_vuln_responses type = {type(hydro_vuln_responses)}")
    print(f"DEBUG: hydro_geo_responses type = {type(hydro_geo_responses)}")
    
    doc.add_heading('1. Executive Summary', 1)
    
    # Initialize summary text
    summary_parts = []
    print("DEBUG: Initialized summary_parts")
    
    # Introduction
    summary_parts.append(f"This report presents the results of a comprehensive hydrologic and hydrogeomorphologic risk assessment for the {project_name} project area, conducted using the USDA Forest Service CULVERT Web Application. The analysis evaluates the vulnerability of road-stream crossing infrastructure to flooding and geomorphologic hazards.")
    print("DEBUG: Added introduction")
    
    # Watershed Delineation Summary
    print("DEBUG: Starting watershed delineation section")
    if ws_deln_responses:
        print(f"DEBUG: ws_deln_responses keys: {list(ws_deln_responses.keys()) if ws_deln_responses else 'None'}")
        hydro_enforcement = ws_deln_responses.get('hydroEnforcementSelect', '')
        print(f"DEBUG: hydro_enforcement = {hydro_enforcement}")
        if hydro_enforcement == 'hydroenf_required':
            summary_parts.append("Watershed delineation was performed employing hydro-enforcement to ensure hydrologic connectivity among watersheds, addressing limitations where elevation data cannot accurately represent drainage systems beneath road infrastructure.")
            print("DEBUG: Added hydro-enforcement required text")
        else:
            summary_parts.append("Watershed delineation was performed without employing hydro-enforcement, utilizing direct digital elevation model processing for drainage boundary determination.")
            print("DEBUG: Added hydro-enforcement not required text")
    else:
        print("DEBUG: No watershed delineation responses available")
    
    # Hydrologic Vulnerability Assessment Summary
    print("DEBUG: Starting hydrologic vulnerability section")
    max_rp = None
    if hydro_vuln_responses:
        print(f"DEBUG: hydro_vuln_responses keys: {list(hydro_vuln_responses.keys()) if hydro_vuln_responses else 'None'}")
        hydro_methods = []
        if hydro_vuln_responses.get('QregionalFrequencyCheckbox') == 'checked':
            hydro_methods.append('Regional Frequency Analysis (RFA)')
            print("DEBUG: Added RFA method")
        if hydro_vuln_responses.get('rationalMethodCheckbox') == 'checked':
            hydro_methods.append('Rational Method (RM)')
            print("DEBUG: Added RM method")
        if hydro_vuln_responses.get('graphPeakCheckbox') == 'checked':
            hydro_methods.append('Graphical Peak Discharge Method (GPDM)')
            print("DEBUG: Added GPDM method")
        
        print(f"DEBUG: hydro_methods = {hydro_methods}")
        
        if hydro_methods:
            methods_text = ', '.join(hydro_methods)
            summary_parts.append(f"The hydrologic vulnerability assessment employed {len(hydro_methods)} analytical methods: {methods_text}.")
            print("DEBUG: Added hydrologic methods text")
            
            # Get maximum return period
            return_periods = hydro_vuln_responses.get('rpListSelect[]', [])
            print(f"DEBUG: return_periods = {return_periods}, type = {type(return_periods)}")
            try:
                if isinstance(return_periods, list) and return_periods:
                    # Handle list case
                    max_rp = max([int(rp) for rp in return_periods])
                    summary_parts.append(f"Analysis was conducted for multiple return periods up to {max_rp} years.")
                    print(f"DEBUG: Set max_rp from list = {max_rp}")
                elif isinstance(return_periods, str) and return_periods:
                    # Handle string case (comma-separated values)
                    if ',' in return_periods:
                        # Split comma-separated string and convert to integers
                        rp_list = [int(rp.strip()) for rp in return_periods.split(',')]
                        max_rp = max(rp_list)
                        summary_parts.append(f"Analysis was conducted for multiple return periods up to {max_rp} years.")
                        print(f"DEBUG: Set max_rp from comma-separated string = {max_rp}")
                    else:
                        # Single value as string
                        max_rp = int(return_periods)
                        summary_parts.append(f"Analysis was conducted for return periods up to {max_rp} years.")
                        print(f"DEBUG: Set max_rp from single string value = {max_rp}")
                elif return_periods:
                    # Try to convert directly to int as fallback
                    max_rp = int(return_periods)
                    summary_parts.append(f"Analysis was conducted for return periods up to {max_rp} years.")
                    print(f"DEBUG: Set max_rp from direct conversion = {max_rp}")
                else:
                    max_rp = None
                    print("DEBUG: No return periods found")
            except Exception as e:
                print(f"DEBUG: Error processing return periods: {e}")
                print(f"DEBUG: return_periods value: '{return_periods}'")
                max_rp = None
    else:
        print("DEBUG: No hydrologic vulnerability responses available")
    
    # Hydrogeomorphologic Vulnerability Assessment Summary
    print("DEBUG: Starting hydrogeomorphologic vulnerability section")
    if hydro_geo_responses:
        print(f"DEBUG: hydro_geo_responses keys: {list(hydro_geo_responses.keys()) if hydro_geo_responses else 'None'}")
        hydrogeo_methods = []
        if hydro_geo_responses.get('sbevaEnabled') == 'true':
            hydrogeo_methods.append('Streambank Erosion Vulnerability Assessment (SBEVA)')
            print("DEBUG: Added SBEVA method")
        if hydro_geo_responses.get('rusleEnabled') == 'true':
            hydrogeo_methods.append('Revised Universal Soil Loss Equation (RUSLE)')
            print("DEBUG: Added RUSLE method")
        if hydro_geo_responses.get('weppEnabled') == 'true':
            hydrogeo_methods.append('Water Erosion Prediction Project (WEPP)')
            print("DEBUG: Added WEPP method")
        if hydro_geo_responses.get('wdfmEnabled') == 'true':
            hydrogeo_methods.append('Watershed Debris Flow Model (WDFM)')
            print("DEBUG: Added WDFM method")
        
        print(f"DEBUG: hydrogeo_methods = {hydrogeo_methods}")
        
        if hydrogeo_methods:
            methods_text = ', '.join(hydrogeo_methods)
            summary_parts.append(f"The hydrogeomorphologic vulnerability assessment integrated {len(hydrogeo_methods)} analytical approaches: {methods_text}.")
            print("DEBUG: Added hydrogeomorphologic methods text")
    else:
        print("DEBUG: No hydrogeomorphologic vulnerability responses available")
    
    # Hydrologic Vulnerability Results Summary
    print("DEBUG: Starting hydrologic vulnerability results section")
    # Hydrologic Vulnerability Results Summary
    if hydro_vuln_responses:
        try:
            # Check for result files from all methods
            rf_path = os.path.join(user_outputs_dir, 'hydro_vuln', 'RFA_results', 'RF_vuln_results.shp')
            rm_path = os.path.join(user_outputs_dir, 'hydro_vuln', 'RM_results', 'RM_vuln_results.shp')
            gpdm_path = os.path.join(user_outputs_dir, 'hydro_vuln', 'GPDM_results', 'GPDM_vuln_results.shp')
            
            print(f"DEBUG: Checking paths:")
            print(f"DEBUG: rf_path exists: {os.path.exists(rf_path)}")
            print(f"DEBUG: rm_path exists: {os.path.exists(rm_path)}")
            print(f"DEBUG: gpdm_path exists: {os.path.exists(gpdm_path)}")
            
            # Dictionary to store results from each method
            method_results = {}
            total_crossings = 0
            bridges = 0
            culverts = 0
            
            # Process Regional Frequency Analysis results
            if os.path.exists(rf_path) and hydro_vuln_responses.get('QregionalFrequencyCheckbox') == 'checked':
                print("DEBUG: Reading RFA results file")
                rf_gdf = gpd.read_file(rf_path)
                if not rf_gdf.empty:
                    print(f"DEBUG: RFA results shape: {rf_gdf.shape}")
                    print(f"DEBUG: RFA columns: {list(rf_gdf.columns)}")
                    
                    # Get basic counts (only once from the first file)
                    if total_crossings == 0:
                        total_crossings = len(rf_gdf['Point_ID'].unique())
                        if 'PourSha' in rf_gdf.columns:
                            bridges = len(rf_gdf[rf_gdf['PourSha'] == 'Bridge'])
                            culverts = total_crossings - bridges
                    
                    method_results['RFA'] = rf_gdf
            
            # Process Rational Method results
            if os.path.exists(rm_path) and hydro_vuln_responses.get('rationalMethodCheckbox') == 'checked':
                print("DEBUG: Reading RM results file")
                rm_gdf = gpd.read_file(rm_path)
                if not rm_gdf.empty:
                    print(f"DEBUG: RM results shape: {rm_gdf.shape}")
                    print(f"DEBUG: RM columns: {list(rm_gdf.columns)}")
                    
                    # Get basic counts if not already set
                    if total_crossings == 0:
                        total_crossings = len(rm_gdf['Point_ID'].unique())
                        if 'PourSha' in rm_gdf.columns:
                            bridges = len(rm_gdf[rm_gdf['PourSha'] == 'Bridge'])
                            culverts = total_crossings - bridges
                    
                    method_results['RM'] = rm_gdf
            
            # Process GPDM results
            if os.path.exists(gpdm_path) and hydro_vuln_responses.get('graphPeakCheckbox') == 'checked':
                print("DEBUG: Reading GPDM results file")
                gpdm_gdf = gpd.read_file(gpdm_path)
                if not gpdm_gdf.empty:
                    print(f"DEBUG: GPDM results shape: {gpdm_gdf.shape}")
                    print(f"DEBUG: GPDM columns: {list(gpdm_gdf.columns)}")
                    
                    # Get basic counts if not already set
                    if total_crossings == 0:
                        total_crossings = len(gpdm_gdf['Point_ID'].unique())
                        if 'PourSha' in gpdm_gdf.columns:
                            bridges = len(gpdm_gdf[gpdm_gdf['PourSha'] == 'Bridge'])
                            culverts = total_crossings - bridges
                    
                    method_results['GPDM'] = gpdm_gdf
            
            # Add crossing counts summary
            if total_crossings > 0:
                summary_parts.append(f"The analysis evaluated {total_crossings} road-stream crossings, comprising {culverts} culverts and {bridges} bridges.")
                print(f"DEBUG: total_crossings = {total_crossings}, culverts = {culverts}, bridges = {bridges}")
            
            # Calculate vulnerability for each method
            if method_results and max_rp:
                vuln_col = f'{max_rp}yrVuln'
                print(f"DEBUG: Looking for vulnerability column: {vuln_col}")
                
                vulnerability_results = []
                
                for method_name, gdf in method_results.items():
                    if vuln_col in gdf.columns and culverts > 0:
                        # Exclude bridges from vulnerability calculation
                        culvert_data = gdf[gdf['PourSha'] != 'Bridge'] if 'PourSha' in gdf.columns else gdf
                        
                        if not culvert_data.empty:
                            vulnerable_culverts = len(culvert_data[culvert_data[vuln_col] == 'Vulnerable'])
                            vuln_percentage = (vulnerable_culverts / culverts * 100)
                            
                            method_full_name = {
                                'RFA': 'Regional Frequency Analysis',
                                'RM': 'Rational Method', 
                                'GPDM': 'Graphical Peak Discharge Method'
                            }.get(method_name, method_name)
                            
                            vulnerability_results.append(f"{method_full_name} identified {vulnerable_culverts} culverts ({vuln_percentage:.1f}%) as vulnerable to the {max_rp}-year flood event")
                            print(f"DEBUG: {method_name} - vulnerable_culverts = {vulnerable_culverts}, percentage = {vuln_percentage:.1f}%")
                
                if vulnerability_results:
                    if len(vulnerability_results) == 1:
                        summary_parts.append(f"Hydrologic vulnerability assessment: {vulnerability_results[0]}.")
                    else:
                        summary_parts.append(f"Hydrologic vulnerability assessment results: {'; '.join(vulnerability_results)}.")
                    print(f"DEBUG: Added vulnerability results for {len(vulnerability_results)} methods")
        
        except Exception as e:
            print(f"DEBUG: Error reading hydrologic vulnerability results: {e}")
            import traceback
            traceback.print_exc()
    
    # Hydrogeomorphologic Vulnerability Results Summary
    print("DEBUG: Starting hydrogeomorphologic vulnerability results section")
    if hydro_geo_responses:
        try:
            hydrogeo_results = []
            
            # SBEVA results
            if hydro_geo_responses.get('sbevaEnabled') == 'true':
                print("DEBUG: Processing SBEVA results")
                sbeva_path = os.path.join(user_outputs_dir, 'hydrogeo_vuln', 'sbeva', 'sbeva_final_output_watershed_polygon.shp')
                print(f"DEBUG: SBEVA path exists: {os.path.exists(sbeva_path)}")
                if os.path.exists(sbeva_path):
                    sbeva_gdf = gpd.read_file(sbeva_path)
                    print(f"DEBUG: SBEVA data shape: {sbeva_gdf.shape}")
                    print(f"DEBUG: SBEVA columns: {list(sbeva_gdf.columns)}")
                    if not sbeva_gdf.empty and 'category' in sbeva_gdf.columns:
                        category_counts = sbeva_gdf['category'].value_counts()
                        print(f"DEBUG: SBEVA category counts: {category_counts}")
                        total = len(sbeva_gdf)
                        if total > 0:
                            high_very_high = category_counts.get('High', 0) + category_counts.get('Very High', 0)
                            high_percentage = (high_very_high / total * 100)
                            hydrogeo_results.append(f"SBEVA analysis shows {high_percentage:.1f}% of watersheds with high to very high streambank erosion vulnerability")
                            print(f"DEBUG: Added SBEVA result: {high_percentage:.1f}%")
            
            # RUSLE results
            if hydro_geo_responses.get('rusleEnabled') == 'true':
                print("DEBUG: Processing RUSLE results")
                rusle_path = os.path.join(user_outputs_dir, 'hydrogeo_vuln', 'rusle', 'rusle_watersheds_with_erosion.shp')
                print(f"DEBUG: RUSLE path exists: {os.path.exists(rusle_path)}")
                if os.path.exists(rusle_path):
                    rusle_gdf = gpd.read_file(rusle_path)
                    print(f"DEBUG: RUSLE data shape: {rusle_gdf.shape}")
                    print(f"DEBUG: RUSLE columns: {list(rusle_gdf.columns)}")
                    if not rusle_gdf.empty and 'category' in rusle_gdf.columns:
                        category_counts = rusle_gdf['category'].value_counts()
                        print(f"DEBUG: RUSLE category counts: {category_counts}")
                        total = len(rusle_gdf)
                        if total > 0:
                            high_very_high = category_counts.get('High', 0) + category_counts.get('Very High', 0)
                            high_percentage = (high_very_high / total * 100)
                            hydrogeo_results.append(f"RUSLE analysis indicates {high_percentage:.1f}% of watersheds with high to very high soil erosion risk")
                            print(f"DEBUG: Added RUSLE result: {high_percentage:.1f}%")
            
            # WDFM results
            if hydro_geo_responses.get('wdfmEnabled') == 'true':
                print("DEBUG: Processing WDFM results")
                wdfm_path = os.path.join(user_outputs_dir, 'hydrogeo_vuln', 'wdfm', 'wdfm_final_output_watershed_polygon.shp')
                print(f"DEBUG: WDFM path exists: {os.path.exists(wdfm_path)}")
                if os.path.exists(wdfm_path):
                    wdfm_gdf = gpd.read_file(wdfm_path)
                    print(f"DEBUG: WDFM data shape: {wdfm_gdf.shape}")
                    print(f"DEBUG: WDFM columns: {list(wdfm_gdf.columns)}")
                    if not wdfm_gdf.empty and 'category' in wdfm_gdf.columns:
                        category_counts = wdfm_gdf['category'].value_counts()
                        print(f"DEBUG: WDFM category counts: {category_counts}")
                        total = len(wdfm_gdf)
                        if total > 0:
                            high_very_high = category_counts.get('High', 0) + category_counts.get('Very High', 0)
                            high_percentage = (high_very_high / total * 100)
                            hydrogeo_results.append(f"WDFM analysis reveals {high_percentage:.1f}% of watersheds with high to very high debris flow potential")
                            print(f"DEBUG: Added WDFM result: {high_percentage:.1f}%")
            # EHVI
            if hydrogeo_results:
                summary_parts.append("Hydrogeomorphologic vulnerability assessment results: " + "; ".join(hydrogeo_results) + ".")
                
                # Add EHVI results if multiple methods were used
                enabled_hydrogeo_methods = []
                if hydro_geo_responses.get('sbevaEnabled') == 'true':
                    enabled_hydrogeo_methods.append('SBEVA')
                if hydro_geo_responses.get('rusleEnabled') == 'true':
                    enabled_hydrogeo_methods.append('RUSLE')
                if hydro_geo_responses.get('weppEnabled') == 'true':
                    enabled_hydrogeo_methods.append('WEPP')
                if hydro_geo_responses.get('wdfmEnabled') == 'true':
                    enabled_hydrogeo_methods.append('WDFM')
                
                if len(enabled_hydrogeo_methods) > 1:
                    try:
                        ehvi_path = os.path.join(user_outputs_dir, 'hydrogeo_vuln', 'ehvi', 'ehvi_final_output_watershed_polygon.shp')
                        print(f"DEBUG: EHVI path exists: {os.path.exists(ehvi_path)}")
                        if os.path.exists(ehvi_path):
                            ehvi_gdf = gpd.read_file(ehvi_path)
                            print(f"DEBUG: EHVI data shape: {ehvi_gdf.shape}")
                            if not ehvi_gdf.empty and 'category' in ehvi_gdf.columns:
                                category_counts = ehvi_gdf['category'].value_counts()
                                print(f"DEBUG: EHVI category counts: {category_counts}")
                                total = len(ehvi_gdf)
                                if total > 0:
                                    high_very_high = category_counts.get('High', 0) + category_counts.get('Very High', 0)
                                    high_percentage = (high_very_high / total * 100)
                                    methods_text = ', '.join(enabled_hydrogeo_methods)
                                    summary_parts.append(f"The Ensemble Hydrogeomorphologic Vulnerability Index (EHVI) integrating {methods_text} shows {high_percentage:.1f}% of watersheds with high to very high composite vulnerability.")
                                    print(f"DEBUG: Added EHVI result: {high_percentage:.1f}%")
                    except Exception as e:
                        print(f"DEBUG: Error reading EHVI results: {e}")
        
        except Exception as e:
            print(f"DEBUG: Error reading hydrogeomorphologic vulnerability results: {e}")
            import traceback
            traceback.print_exc()
    
    # Conclusion
    summary_parts.append("Results provide critical information for infrastructure planning, maintenance prioritization, and risk management decisions, enabling evidence-based approaches to transportation system resilience under changing environmental conditions.")
    print("DEBUG: Added conclusion")
    
    # Join all parts and add to document
    print(f"DEBUG: Total summary parts: {len(summary_parts)}")
    summary_text = "\n\n".join(summary_parts)
    doc.add_paragraph(summary_text)
    doc.add_page_break()
    print("DEBUG: Completed add_executive_summary function")
    
def add_introduction(doc):
    """Add introduction section"""
    doc.add_heading('2. Introduction', 1)
    
    intro_text = """
The CULVERT Web Application represents a comprehensive analytical framework for evaluating hydrologic and hydrogeomorphologic risks to road-stream crossing infrastructure. This assessment is particularly critical for forest transportation systems where culverts, bridges, fords, etc., are exposed to extreme precipitation induced hydrologic events (floods) and geomorphologic processes that can compromise their structural integrity and operational capacity.

Transportation infrastructure in forested environments faces increasing challenges from climate variability, changing precipitation patterns, and evolving watershed conditions. The CULVERT framework addresses these challenges by integrating and automating multiple statistical, empirical, qualitative and analytical methods, automated data acquisition systems, and modern risk assessment approaches.

This analysis provides stakeholders with quantitative and qualitative pre-field visit risk assessments that support evidence-based decision-making for infrastructure investments, maintenance scheduling, and adaptive management strategies. The methodology emphasizes scientific rigor while maintaining practical applicability for operational forest management contexts.
"""
    
    doc.add_paragraph(intro_text)
def add_study_area_overview(doc, generated_plots, project_name):
    """Add study area overview section with images"""
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import os
    
    doc.add_heading('3. Study Area Overview', 1)
    
    study_area_text = f"""
    --- ADD YOUR OWN STUDY AREA CHARACTERISTICS --- 
    
    OR EDIT the sample text below copied from Amatya, Devendra M. and Trettin, Carl C. (2019) "Long-Term Ecohydrologic Monitoring: A Case Study from the Santee Experimental Forest, South Carolina," Journal of South Carolina Water Resources: Vol. 6 : Iss. 1 , Article 5. DOI: https://doi.org/10.34068/JSCWR.06.05"
    
    The SEF was established in 1937 by the USDA Forest Service with a mission of silvicultural research, environmental monitoring, and demonstration, and educational activities in support of sustainable forest management practices of coastal plain forests, such as those within the Francis Marion National Forest (FMNF) near Huger, South Carolina, 50 km northwest of Charleston. The SEF (33˚ 08' 15" N, 79˚ 49' 0" W) is located within the headwaters of Huger Creek, a tributary of the East Branch of the Cooper River that drains into Charleston Harbor (Figure 1). 
    
    In order to study the effects of silvicultural practices on hydrology and water quality, gauged watersheds were established beginning in November 1963 with WS77, a first-order watershed of 155 ha, and WS78, a third-order watershed of 5,240 ha (Amatya and Trettin, 2007; Amatya et al., 2015). A second-order watershed (WS79) of 500 ha was gauged in 1966, followed in 1968 by another first-order watershed (WS80) of 206 ha (reduced to 160 ha in late 2001) as a pair to WS77. The monitoring was discontinued in May 1982, resumed in November 1989 soon after the passage of Hurricane Hugo, and continues to the present day. 
    
    The forest was heavily impacted by the hurricane in 1989 (Hook et al., 1991), and its current vegetation consists of pine and pine mixed with hardwood stands that have been vigorously re-growing since this tropical storm. Soils in SEF are predominantly Alfisols and Ultisols (SCS 1980), primarily somewhat-poorly to poorly drained sandy loams with clayey subsoils with high surface water retention capacity and low permeability. The climate of the site is subtropical with long, hot summers followed by short, warm, and humid winters, with an average annual temperature and potential evapotranspiration (PET) of 18.3˚C and 1135 mm, respectively, as well as an average annual precipitation of 1370 mm (Dai et al., 2013)."

Figure 3.1 presents the comprehensive study area overview for the area of interest, providing essential spatial context for the subsequent vulnerability analyses. The map integrates multiple geospatial datasets including high-resolution digital elevation models, watershed boundary delineations, transportation network infrastructure, and road-stream crossing locations. 

The topographic representation utilizes hillshade relief visualization overlaid on elevation data to highlight terrain characteristics that influence hydrologic flow patterns and geomorphologic processes. Watershed boundaries define the contributing drainage areas for each road-stream crossing, establishing the spatial framework for all subsequent hydrologic calculations. The road network represents the transportation infrastructure system requiring protection, while road-stream crossing locations identify specific points of vulnerability where infrastructure intersects natural drainage systems.

The study area encompasses diverse topographic conditions and drainage characteristics that create varying levels of flood risk and geomorphologic hazard exposure across different crossing locations. This spatial heterogeneity necessitates location-specific vulnerability assessments and adaptive management strategies tailored to local environmental conditions.
"""
    
    doc.add_paragraph(study_area_text.strip())
    
    # Add study area image if available
    if 'study-area' in generated_plots and os.path.exists(generated_plots['study-area']):
        try:
            # Add image with proper sizing (6.5 inches width maintains aspect ratio)
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(generated_plots['study-area'], width=Inches(6.5))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        except Exception as e:
            print(f"Error adding study area image: {e}")
            # Add placeholder text if image fails
            doc.add_paragraph("[Study Area Map - Image not available]").alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[Study Area Map - Image not available]").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add figure caption immediately after image
    caption = doc.add_paragraph()
    caption_run = caption.add_run('Figure 3.1: ')
    caption_run.bold = True
    caption_run.font.size = Pt(10)
    
    caption_text = caption.add_run(f'Study area overview for area of interest showing elevation (m), region boundary, road network infrastructure, and road-stream crossing locations (if available).')
    caption_text.font.size = Pt(10)
    caption_text.italic = True
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
def add_data_summary_table(doc):
    """Add data summary table"""
    doc.add_heading('4. Data Summary', 1)
    
    doc.add_paragraph('Table 4.1. Summary of Datasets used in the CULVERT Web-Application')
    
    # Create table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # Add headers
    headers = ['Sl No.', 'Variable Name', 'Method of Data Ingestion', 'Source & Description', 'Data Availability', 'Application in CULVERT']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Make headers bold
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data rows (abbreviated for space)
    data_rows = [
        ['1', 'Region Boundary Polygon', 'User Upload', 'User-dependent source (zipped Shapefile)', 'Mandatory', 'Watershed delineation, spatial analysis'],
        ['2', 'Digital Elevation Model', 'User Upload', 'User-dependent source (GeoTIFF)-available here https://datagateway.nrcs.usda.gov/GDGOrder.aspx', 'Mandatory', 'Flow analysis, slope calculation, watershed delineation'],
        ['3', 'Pour Point Locations and drainage sturcture dimension', 'User Upload', 'zipped Point Shapefile with attributes', 'Optional', 'Discharge capacity estimation & Infrastructure vulnerability assessment'],
        ['4', 'Road Network Data', 'User Upload/API', 'zipped Shapefile or OpenStreetMap API (https://www.openstreetmap.org/#map=5/38.01/-95.84)', 'Optional', 'Watershed delineation and road-stream intersection analysis'],
        ['5', 'Land Cover Data', 'API', 'NLCD 2021 (https://www.mrlc.gov/data-services-page)', 'Automatic', 'Runoff coefficient estimation, Curve Number estimation, SBEVA and WDFM vulnerability assessment'],
        ['6', 'Soil Data', 'Pre-processed/Database', 'gSSURGO Database (https://www.nrcs.usda.gov/resources/data-and-reports/gridded-soil-survey-geographic-gssurgo-database)', 'Automatic', 'Hydrologic soil group classification'],
        ['7', 'Precipitation Data', 'User/API', 'Onsite rain gauge or NOAA Atlas-14 (https://hdsc.nws.noaa.gov/pfds/)', 'Optional/Auto', 'Frequency analysis, runoff calculation'],
        ['8', 'Streamflow Data', 'User Upload', 'Onsite stream gauge', 'Optional', 'Regional frequency analysis'],
        ['9', 'Climate Data', 'API', 'PRISM 30-year normals (https://prism.oregonstate.edu/)', 'Automatic', 'watershed characteristics, SBEVA and WDFM vulnerability assessment'],
        ['10', 'Wetland Cover', 'Pre-processed/Database', 'U.S. Fish & Wildlife Service (https://www.fws.gov/program/national-wetlands-inventory/wetlands-mapper)', 'Automatic', 'Watershed characteristics, region of influence analysis'],
        ['11', 'NDVI', 'Pre-processed/Database', 'USGS eVIIRS (https://www.usgs.gov/centers/eros/science/usgs-eros-archive-vegetation-monitoring-eros-visible-infrared-imaging)', 'Automatic', 'WDFM vulnerability assessment'],
        ['12', 'Geology Data', 'Pre-processed/Database', 'USGS Minearl Resources (https://www.usgs.gov/programs/mineral-resources-program/data)', 'Automatic', 'WDFM vulnerability assessment'],
    ]
    
    for row_data in data_rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = cell_text

def add_results_section(doc, generated_plots, project_name):
    """Add comprehensive results and analysis section with images"""
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import os
    
    doc.add_heading('5. Results and Analysis', 1)
    
    # Add introductory paragraph for the results section
    intro_text = f"""
This section presents the comprehensive analysis results for the {project_name} project, including watershed delineation outcomes, hydrologic vulnerability assessments, and hydrogeomorphologic risk evaluations. The analysis integrates multiple approaches to provide robust infrastructure vulnerability assessments and risk quantification across different temporal and spatial scales. Each subsection includes detailed spatial visualizations, statistical summaries, and technical interpretations to support evidence-based decision-making for transportation infrastructure management.
"""
    doc.add_paragraph(intro_text.strip())
    
    # 5.1 Watershed and Road-Stream Crossing Analysis
    doc.add_heading('5.1 Watershed and Road-Stream Crossing Analysis', 2)
    
    watershed_analysis_text = f"""
    --- ADD YOUR OWN OBSERVATIONS FROM Figure 5.1 --- 
    
    OR EDIT the text below
    
Figure 5.1 presents detailed spatial analysis of watershed characteristics and road-stream crossing infrastructure attributes across the {project_name} study area. The comprehensive visualization includes four key watershed parameters and three infrastructure classification schemes that collectively inform hydrologic vulnerability assessments and infrastructure management decisions.

Watershed characteristics analysis reveals the spatial distribution of drainage areas (hectares), average slope gradients, main channel lengths (meters), and time of concentration values (minutes). These parameters directly influence peak discharge calculations, flood timing, and runoff generation processes. Larger drainage areas typically generate higher peak discharges but may have longer concentration times, while steeper slopes increase runoff velocities and reduce lag times between precipitation and peak flows.

Infrastructure analysis documents the spatial distribution of crossing types (culverts, bridges, fords), construction materials (concrete, steel, wood, composite), and current condition assessments (excellent, good, fair, poor). This information enables prioritization of maintenance activities, replacement planning, and adaptive management strategies based on both hydraulic vulnerability and structural condition.

The spatial clustering of infrastructure attributes reveals systematic patterns in construction approaches and maintenance needs across the transportation network. Areas with older infrastructure or harsh environmental conditions may exhibit concentrated maintenance requirements, while newer installations may demonstrate improved resilience to extreme hydrologic events.
"""
    
    doc.add_paragraph(watershed_analysis_text.strip())
    
    # Add watershed analysis image if available
    if 'watershed-analysis' in generated_plots and os.path.exists(generated_plots['watershed-analysis']):
        try:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(generated_plots['watershed-analysis'], width=Inches(6.5))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        except Exception as e:
            print(f"Error adding watershed analysis image: {e}")
            doc.add_paragraph("[Watershed Analysis Maps - Image not available]").alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[Watershed Analysis Maps - Image not available]").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add figure caption immediately after image
    caption = doc.add_paragraph()
    caption_run = caption.add_run('Figure 5.1: ')
    caption_run.bold = True
    caption_run.font.size = Pt(10)
    
    caption_text = caption.add_run('Comprehensive watershed characteristics and road-stream crossing infrastructure analysis showing spatial distribution of drainage areas, slope gradients, channel lengths, time of concentration values, crossing types, construction materials, and condition assessments.')
    caption_text.font.size = Pt(10)
    caption_text.italic = True
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 5.2 Hydrologic Risk Analysis
    doc.add_heading('5.2 Hydrologic Risk Analysis', 2)
    
    hydro_risk_text = f"""
    --- ADD YOUR OWN OBSERVATIONS FROM Figure 5.2 --- 
    
    OR EDIT the text below
    
Figure 5.2 presents comprehensive hydrologic vulnerability analysis results for the {project_name} project, integrating multiple analytical methods to provide robust flood risk assessments across different return periods and methodological approaches. The analysis includes vulnerability classifications, peak discharge estimates, discharge capacity evaluations, and statistical distributions that collectively inform infrastructure risk management decisions.

The hydrologic vulnerability assessment employs up to three complementary analytical methods: Regional Frequency Analysis (RFA) for ungauged watershed conditions, the Rational Method (RM) for rapid assessment applications, and the Graphical Peak Discharge Method (GPDM) for standardized engineering calculations. Each method provides independent estimates of flood magnitudes for specified return periods, enabling uncertainty quantification and methodological validation.

Vulnerability classifications identify crossings as "Vulnerable" or "Not Vulnerable" based on comparison between estimated peak discharges and calculated infrastructure capacity. Crossings classified as vulnerable require immediate attention for potential capacity upgrades, structural modifications, or enhanced maintenance protocols. The spatial distribution of vulnerable crossings reveals systematic patterns related to watershed size, topographic position, and infrastructure age.

Peak discharge estimates demonstrate the range of flood magnitudes expected across different return periods, with higher return periods generating correspondingly larger discharge values. The methodology comparison reveals convergence or divergence among analytical approaches, providing confidence intervals for design decisions and highlighting areas requiring additional data collection or analysis refinement.

Discharge capacity analysis evaluates the hydraulic performance of existing infrastructure under design flood conditions. The capacity calculations consider geometric constraints, hydraulic efficiency, and structural limitations that may restrict flow conveyance during extreme events. Crossings with insufficient capacity relative to design floods represent priority locations for infrastructure improvements.

The statistical distribution analysis provides comprehensive comparison of discharge estimates across all methods and return periods, revealing central tendencies, variability ranges, and outlier conditions that inform risk assessment uncertainty. Box plots demonstrate the statistical properties of discharge distributions while identifying crossings with exceptional flood risk characteristics.
"""
    
    doc.add_paragraph(hydro_risk_text.strip())
    
    # Add hydrologic risk image if available
    if 'hydro-risk' in generated_plots and os.path.exists(generated_plots['hydro-risk']):
        try:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(generated_plots['hydro-risk'], width=Inches(6.5))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        except Exception as e:
            print(f"Error adding hydrologic risk image: {e}")
            doc.add_paragraph("[Hydrologic Risk Analysis Maps - Image not available]").alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[Hydrologic Risk Analysis Maps - Image not available]").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add figure caption immediately after image
    caption = doc.add_paragraph()
    caption_run = caption.add_run('Figure 5.2: ')
    caption_run.bold = True
    caption_run.font.size = Pt(10)
    
    caption_text = caption.add_run('Comprehensive hydrologic risk analysis showing vulnerability classifications, peak discharge estimates across multiple methods and return periods, discharge capacity evaluations, and statistical distributions of flood magnitudes for infrastructure planning and risk management.')
    caption_text.font.size = Pt(10)
    caption_text.italic = True
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 5.3 Hydrogeomorphologic Risk Analysis  
    doc.add_heading('5.3 Hydrogeomorphologic Risk Analysis', 2)
    
    hydrogeo_risk_text = f"""
    
--- ADD YOUR OWN OBSERVATIONS FROM Figure 5.3 --- 
    
    OR EDIT the text below
    
Figure 5.3 presents comprehensive hydrogeomorphologic vulnerability assessment results for the {project_name} project, integrating multiple erosion and mass wasting analytical frameworks to evaluate long-term threats to transportation infrastructure stability. The analysis encompasses streambank erosion potential, debris flow hazards, soil erosion rates, ensemble vulnerability indices, and statistical distributions that collectively inform geomorphologic risk management strategies.

The Streambank Erosion Vulnerability Assessment (SBEVA) evaluates lateral channel erosion processes that threaten infrastructure through bank undermining, channel widening, and flow pattern alteration. SBEVA integrates climatic forcing variables (precipitation intensity, temperature, solar radiation), soil properties (available water storage, drainage characteristics, hydrologic soil groups), topographic factors (slope gradients), and land cover conditions to generate composite vulnerability scores. High SBEVA scores indicate locations where streambank erosion processes pose significant threats to crossing stability and long-term performance.

The Watershed Debris Flow Model (WDFM) assesses mass wasting hazards that can block or damage infrastructure through debris delivery from upslope source areas. WDFM incorporates slope stability factors, soil properties, geological characteristics, vegetation cover, and climatic triggers to identify areas susceptible to debris flow initiation and transport. High WDFM scores highlight crossings vulnerable to debris blockage, structural impact damage, and operational disruption during extreme precipitation events.

The Revised Universal Soil Loss Equation (RUSLE) quantifies sheet and rill erosion rates across watershed surfaces, providing estimates of sediment production that may accumulate in channels and impact infrastructure performance. RUSLE calculations integrate rainfall erosivity, soil erodibility, topographic factors, vegetation cover, and conservation practices to estimate annual soil loss rates in kilograms per year. High erosion rates indicate areas generating excessive sediment loads that may require channel maintenance, debris removal, or sediment control measures.

The Ensemble Hydrogeomorphological Vulnerability Index (EHVI) provides integrated risk assessment by combining SBEVA, WDFM, and RUSLE results into unified vulnerability classifications. EHVI scores represent composite risk levels that account for multiple geomorphologic processes threatening infrastructure performance. The ensemble approach reduces individual method limitations while providing comprehensive risk characterization for management decision-making.

Vulnerability distribution statistics reveal the proportion of crossings classified within each risk category (Very High, High, Moderate, Low, Very Low) across all analytical methods. These distributions inform prioritization strategies, budget allocation decisions, and adaptive management planning by quantifying the scale of vulnerability across the transportation network.

Spatial patterns in hydrogeomorphologic vulnerability reflect underlying geological, topographic, and climatic controls on erosion processes. Areas with steep slopes, erosive soils, intense precipitation, or sparse vegetation cover typically exhibit elevated vulnerability scores, while stable terrain with protective vegetation demonstrates lower risk levels.
"""
    
    doc.add_paragraph(hydrogeo_risk_text.strip())
    
    # Add hydrogeomorphologic risk image if available
    if 'hydrogeo-risk' in generated_plots and os.path.exists(generated_plots['hydrogeo-risk']):
        try:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(generated_plots['hydrogeo-risk'], width=Inches(6.5))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        except Exception as e:
            print(f"Error adding hydrogeomorphologic risk image: {e}")
            doc.add_paragraph("[Hydrogeomorphologic Risk Analysis Maps - Image not available]").alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[Hydrogeomorphologic Risk Analysis Maps - Image not available]").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add figure caption immediately after image
    caption = doc.add_paragraph()
    caption_run = caption.add_run('Figure 5.3: ')
    caption_run.bold = True
    caption_run.font.size = Pt(10)
    
    caption_text = caption.add_run('Comprehensive hydrogeomorphologic risk analysis showing streambank erosion vulnerability (SBEVA), debris flow potential (WDFM), soil erosion rates (RUSLE), ensemble vulnerability indices (EHVI), and vulnerability distribution statistics across multiple geomorphologic processes.')
    caption_text.font.size = Pt(10)
    caption_text.italic = True
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 5.4 Integrated Risk Assessment Summary
    doc.add_heading('5.4 Integrated Risk Assessment Summary', 2)
    
    summary_text = f"""
--- ADD YOUR OWN SUMMARY FROM THE RESULTS --- 
    
    OR EDIT the text below
The comprehensive risk assessment for the {project_name} project reveals complex spatial patterns of infrastructure vulnerability that reflect the interaction between hydrologic extremes, geomorphologic processes, and infrastructure characteristics. The multi-method analytical approach provides robust uncertainty quantification while identifying priority locations for management intervention.

Key findings from the integrated assessment include:

• Spatial heterogeneity in flood risk across the transportation network, with vulnerability concentrations in areas of large drainage areas, steep terrain, and aging infrastructure
• Convergence among hydrologic methods in most areas, providing confidence in discharge estimates, with divergence highlighting areas requiring additional data collection or analysis refinement  
• Systematic patterns in hydrogeomorphologic vulnerability related to geological, topographic, and vegetation controls on erosion processes
• Infrastructure capacity limitations that constrain flood conveyance at multiple crossing locations, particularly for higher return period events
• Correlation between hydrologic and hydrogeomorphologic vulnerabilities in areas of steep terrain and intense precipitation, suggesting compound risk scenarios requiring comprehensive management strategies

The vulnerability assessment results support evidence-based prioritization of infrastructure improvements, maintenance scheduling, and adaptive management planning. Crossings identified as vulnerable under multiple analytical approaches represent highest priority locations for immediate intervention, while areas showing emerging risks may benefit from proactive management measures.

Uncertainty quantification through multiple methods provides confidence intervals for design decisions while highlighting data limitations that may benefit from additional monitoring or analysis. The statistical distributions reveal the range of expected conditions and help establish appropriate safety factors for infrastructure design and operation.
"""
    
    doc.add_paragraph(summary_text.strip())
    
    # 5.5 Conclusions and Recommendations
    doc.add_heading('5.5 Conclusions', 2)
    
    conclusions_text = f"""
--- ADD YOUR OWN CONCLUSION FROM THE RESULTS --- 
    
    OR EDIT the text below
The comprehensive risk assessment for the {project_name} project provides critical information for sustainable transportation infrastructure management under changing environmental conditions. The analysis integrates state-of-the-art hydrologic and hydrogeomorphologic assessment methods to deliver quantitative risk estimates that support evidence-based decision-making.

Primary conclusions include:

1. Vulnerability Identification: The analysis successfully identified specific road-stream crossings requiring immediate attention due to inadequate hydraulic capacity or high geomorphologic risk exposure.

2. Risk Quantification: Multiple analytical methods provide robust uncertainty bounds for flood risk estimates, enabling appropriate safety factor selection for design and operation decisions.

3. Spatial Risk Patterns: Systematic vulnerability patterns reflect underlying environmental controls, allowing targeted management strategies adapted to local conditions.

4. Infrastructure Performance: Existing crossing capacities demonstrate varying adequacy levels across different return periods, with clear priorities for capacity enhancement or replacement.

5. Methodological Validation: Convergence among independent analytical approaches provides confidence in risk estimates, while divergence highlights areas requiring additional investigation.

Priority Recommendations:

Immediate Actions (0-2 years):
• Prioritize maintenance and potential replacement of crossings identified as hydrologically vulnerable across multiple methods
• Implement emergency preparedness protocols for crossings in high-risk areas during extreme weather events  
• Conduct detailed hydraulic analyses for crossings showing capacity limitations under design return periods
• Establish monitoring protocols for erosion and sedimentation patterns in high RUSLE erosion rate areas

Medium-term Planning (2-5 years):
• Develop adaptive management strategies for crossings in areas with high hydrogeomorphologic vulnerability (EHVI scores)
• Consider climate adaptation measures for infrastructure in areas showing non-stationary flood frequency trends
• Implement debris management and channel maintenance programs for crossings with elevated WDFM scores
• Plan capacity upgrades for crossings with insufficient hydraulic performance under projected future conditions

Long-term Strategic Planning (5+ years):
• Integrate risk assessment results with broader transportation system planning and forest management objectives
• Develop regional infrastructure resilience strategies based on vulnerability clustering patterns  
• Establish adaptive management frameworks that can respond to changing environmental conditions
• Consider relocating or redesigning crossings in areas with persistently high vulnerability across multiple hazard types

Data and Analysis Enhancement:
• Install stream gauging equipment in ungauged watersheds showing high vulnerability or methodological uncertainty
• Collect additional precipitation data in areas with limited climate station coverage
• Conduct periodic vulnerability reassessments to track changing risk patterns over time
• Validate model predictions through post-event damage assessment and infrastructure performance monitoring

These recommendations should be integrated with broader forest management objectives, budgetary constraints, and regulatory requirements to ensure sustainable infrastructure performance under evolving environmental conditions. The quantitative risk assessment framework established through this analysis provides a foundation for adaptive management approaches that can respond effectively to changing hydrologic and geomorphologic conditions.
"""
    
    doc.add_paragraph(conclusions_text.strip())
    
    # 5.6 Acknowledgment
    doc.add_heading('5.5 Acknowledgement', 2)
    
    summary_text = f"""
--- ADD YOUR OWN ACKNOWLEDGEMENT --- 
    
    OR EDIT the text below:
    
    We would like to gratefully acknowledge USDA-FS for the open-source CULVERT Web-application and datasets provided by USDA, NOAA, USGS, NRCS, U.S. Fish & Wildlife Service, and Openstreet Map. 
"""
    
    doc.add_paragraph(summary_text.strip())
    
def add_appendices(doc, ws_deln_responses, hydro_vuln_responses, hydro_geo_responses):
    """Add comprehensive appendices with technical details and parameters"""
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import ast
    
    # ========================================
    # APPENDIX A: ANALYSIS PARAMETERS
    # ========================================
    doc.add_heading('Appendix A: Analysis Parameters', 1)
    
    appendix_intro = """
This appendix provides comprehensive documentation of all analysis parameters, user selections, and configuration settings employed in the CULVERT assessment. These parameters directly influence the analytical results and should be considered when interpreting findings or replicating the analysis. Parameter values reflect user-specified settings, default system values, or automatically determined optimal configurations based on input data characteristics.
"""
    doc.add_paragraph(appendix_intro.strip())
    
    # ========================================
    # A.1 WATERSHED DELINEATION PARAMETERS
    # ========================================
    if ws_deln_responses:
        doc.add_heading('A.1 Watershed Delineation Parameters', 2)
        
        ws_intro = """
Watershed delineation parameters control the spatial analysis procedures used to define drainage boundaries, stream networks, and pour point locations. These parameters significantly influence the accuracy of hydrologic modeling results and should be selected based on terrain characteristics, data resolution, and analysis objectives.
"""
        doc.add_paragraph(ws_intro.strip())
        
        # Create comprehensive parameter table
        ws_table = doc.add_table(rows=1, cols=3)
        ws_table.style = 'Table Grid'
        
        # Set column widths for better readability
        ws_table.columns[0].width = Inches(2.5)  # Parameter Name
        ws_table.columns[1].width = Inches(1.0)  # Value
        ws_table.columns[2].width = Inches(3.0)  # Description
        
        # Add headers
        ws_headers = ['Parameter', 'Value', 'Description']
        ws_header_cells = ws_table.rows[0].cells
        for i, header in enumerate(ws_headers):
            ws_header_cells[i].text = header
            for paragraph in ws_header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
        
        # Custom parameter processing with specific logic
        processed_params = {}
        
        # Process each parameter with custom logic
        for key, value in ws_deln_responses.items():
            if key == 'boundaryShapefile':
                processed_params['Region Boundary Shapefile'] = ('Uploaded', 'User-provided polygon defining the analysis extent')
            
            elif key == 'demRaster':
                processed_params['Digital Elevation Model Data'] = ('Uploaded', 'User-provided raster dataset for terrain analysis')
            
            elif key == 'RoadDataSelect':
                if value == '{}' or not value or value == 'null':
                    processed_params['Road Data Source'] = ('API', 'Road Data acquired from Open-street Map API https://www.openstreetmap.org')
                else:
                    processed_params['Road Data Source'] = ('Road Data Uploaded by user', 'Road Data acquired from user-provided dataset')
            
            elif key == 'pourPointDataSelect':
                if value == 'both':
                    processed_params['Pour Point Data Source'] = ('Both RSC and Gauging Station Data', 'Method for defining crossing locations')
                elif value == 'pour_pt_NA':
                    processed_params['Pour Point Data Source'] = ('No data available', 'Method for defining crossing locations')
                elif value == 'culvert':
                    processed_params['Pour Point Data Source'] = ('RSC data available only', 'Method for defining crossing locations')
                elif value == 'gauging':
                    processed_params['Pour Point Data Source'] = ('Only gauging station data', 'Method for defining crossing locations')
                else:
                    processed_params['Pour Point Data Source'] = (format_parameter_value(value), 'Method for defining crossing locations')
            
            elif key == 'hydroEnforcementSelect':
                if value == 'hydroenf_required':
                    processed_params['Hydro-enforcement Option'] = ('Hydro-enforced', 'Terrain modification approach for improved flow routing')
                else:
                    processed_params['Hydro-enforcement Option'] = ('Not Hydro-enforced', 'Terrain modification approach for improved flow routing')
            
            elif key == 'roadFillDemByM':
                processed_params['Road Fill DEM Height (m)'] = (format_parameter_value(value), 'Vertical adjustment applied to road elevations')
            
            elif key == 'roadFillDemBufferM':
                processed_params['Road Fill Buffer Distance (m)'] = (format_parameter_value(value), 'Lateral extent of road elevation modifications')
            
            elif key == 'breaklineOffsetM':
                processed_params['Breakline Offset Distance (m)'] = (format_parameter_value(value), 'Perpendicular distance for breakline placement')
            
            elif key == 'breaklineBurnDemByM':
                processed_params['Breakline Burn Depth (m)'] = (format_parameter_value(value), 'Vertical depression applied along breaklines')
            
            elif key == 'breaklineBurnDemBufferM':
                processed_params['Breakline Burn Buffer (m)'] = (format_parameter_value(value), 'Lateral extent of breakline modifications')
            
            elif key == 'flowAccumThreshold':
                processed_params['Flow Accumulation Threshold'] = (format_parameter_value(value), 'Minimum contributing area for stream initiation')
            
            elif key == 'pourPointSnapDistanceM':
                processed_params['Pour Point Snap Distance (m)'] = (format_parameter_value(value), 'Maximum distance for pour point alignment to streams')
            
            elif key == 'filterWatershedMinAreaHa':
                processed_params['Minimum Watershed Area (Ha)'] = (format_parameter_value(value), 'Lower size limit for watershed inclusion')
            
            elif key == 'flagWatershedAreaOutsideBoundaryHa':
                processed_params['Boundary Overage Threshold (Ha)'] = (format_parameter_value(value), 'Maximum allowable watershed area beyond region boundary')
            
            elif key == 'FlagKeepOptionSelect':
                processed_params['Flagged Watershed Handling'] = (format_parameter_value(value), 'Treatment of watersheds exceeding boundary limits')
        
        # Add processed parameter rows to table
        for param_name, (param_value, description) in processed_params.items():
            row_cells = ws_table.add_row().cells
            
            row_cells[0].text = param_name
            row_cells[1].text = param_value
            row_cells[2].text = description
            
            # Format cells
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        
        # Add any additional parameters not processed above
        processed_keys = {'boundaryShapefile', 'demRaster', 'RoadDataSelect', 'RoadData', 'pourPointDataSelect', 'PourData', 
                         'hydroEnforcementSelect', 'roadFillDemByM', 'roadFillDemBufferM', 'breaklineOffsetM', 
                         'breaklineBurnDemByM', 'breaklineBurnDemBufferM', 'flowAccumThreshold', 'pourPointSnapDistanceM', 
                         'filterWatershedMinAreaHa', 'flagWatershedAreaOutsideBoundaryHa', 'FlagKeepOptionSelect'}
        
        additional_params = {k: v for k, v in ws_deln_responses.items() if k not in processed_keys}
        if additional_params:
            doc.add_heading('A.1.1 Additional Watershed Delineation Parameters', 3)
            for key, value in additional_params.items():
                param_para = doc.add_paragraph()
                param_para.add_run(f"• {key.replace('_', ' ').title()}: ").bold = True
                param_para.add_run(format_parameter_value(value))
    
    # ========================================
    # A.2 HYDROLOGIC VULNERABILITY PARAMETERS
    # ========================================
    if hydro_vuln_responses:
        doc.add_heading('A.2 Hydrologic Vulnerability Assessment Parameters', 2)
        
        hydro_intro = """
    Hydrologic vulnerability parameters control the statistical analysis methods, data sources, and calculation procedures used for flood frequency analysis and peak discharge estimation. These parameters determine the robustness of risk assessments and the applicability of results to specific return periods and design scenarios.
    """
        doc.add_paragraph(hydro_intro.strip())
        
        def format_statistical_parameter_value(key, value):
            """Format statistical parameter values with proper display names"""
            
            # Distribution type formatting
            if key in ['QtyofDistSelect', 'PItyofDistSelect']:
                if value == 'GEV':
                    return 'Generalized Extreme Value Distribution'
                elif value == 'Gumbel':
                    return 'Gumbel Distribution'
                elif value == 'Pearson3':
                    return 'Pearson Type III Distribution'
                elif value == 'LogPearson3':
                    return 'Log-Pearson Type III Distribution'
            
            # Homogeneity test formatting
            elif key in ['QhomTestSelect', 'PIhomTestSelect']:
                if value == 'AD':
                    return 'Anderson-Darling Homogeneity Test'
                elif value == 'HW':
                    return 'Hosking-Wallis Homogeneity Test'
                elif value == 'KS':
                    return 'Kolmogorov-Smirnov Test'
            
            # Outlier detection method formatting
            elif key in ['QoutlierMethodSelect', 'PIoutlierMethodSelect']:
                if value == 'None' or value == 'NA':
                    return 'No Outlier Removal'
                elif value == 'Zscore':
                    return 'Z-Score Method'
                elif value == 'IQR':
                    return 'Inter-Quartile Range Method'
                elif value == 'Grubbs':
                    return 'Grubbs Test'
            
            # Z-score threshold formatting
            elif key in ['QzScoreSelect', 'PIzScoreSelect']:
                return f'Z-Score = {value}'
            
            # Parameter estimation method formatting
            elif key in ['QparEstSelect', 'PIparEstSelect']:
                if value == 'MLE':
                    return 'Maximum Likelihood Estimation'
                elif value == 'GMLE':
                    return 'Generalized Maximum Likelihood Estimation'
                elif value == 'Bayesian':
                    return 'Bayesian Parameter Estimation'
                elif value == 'MOM':
                    return 'Method of Moments'
                elif value == 'LMOM':
                    return 'L-Moments Method'
            
            # Stationarity assumption formatting
            elif key in ['QNonStationarySelect', 'PINonStationarySelect']:
                if value == 'Stationarity':
                    return 'Stationary Analysis'
                elif value == 'Non-stationarity':
                    return 'Non-Stationary Analysis'
            
            # Precipitation data source formatting
            elif key == 'PIdataSelect':
                if value == 'NOAAatlas14':
                    return 'NOAA Atlas-14'
                elif value == 'From Watershed Specific Rain Gauges':
                    return 'Watershed Specific Rain Gauges'
                elif value == 'User':
                    return 'User-Provided Data'
                elif value == 'API':
                    return 'API Retrieved Data'
            
            # Runoff coefficient method formatting
            elif key == 'CoeffRunoffSelect':
                if value == 'fromTable':
                    return 'Referring Table'
                elif value == 'backCalculation':
                    return 'Back-Calculation'
                elif value == 'NRCS':
                    return 'NRCS Method'
            
            # Culvert capacity method formatting
            elif key == 'CulvertCapacitySelect':
                if value == 'orifice_flow':
                    return 'Orifice Flow (Simplest Method)'
                elif value == 'inlet_control':
                    return 'Inlet Control (FHWA HDS-5)'
                elif value == 'outlet_control':
                    return 'Outlet Control (FHWA HDS-5)'
                elif value == 'manning_uniform':
                    return "Manning's Uniform Flow"
            
            # Default formatting for any unmatched values
            return format_parameter_value(value)
        
        # A.2.1 General Analysis Configuration
        doc.add_heading('A.2.1 General Analysis Configuration', 3)
        
        # Extract general configuration
        general_config = {}
        if 'Gst_Names' in hydro_vuln_responses:
            watersheds = hydro_vuln_responses['Gst_Names']
            if isinstance(watersheds, list):
                general_config['Analyzed Gauged Watersheds'] = f"{len(watersheds)} watersheds: {', '.join(watersheds)}"
            else:
                general_config['Analyzed Gauged Watersheds'] = str(watersheds)
        
        # Methods enabled
        methods_enabled = []
        if hydro_vuln_responses.get('QregionalFrequencyCheckbox') == 'checked':
            methods_enabled.append('Regional Frequency Analysis (RFA)')
        if hydro_vuln_responses.get('rationalMethodCheckbox') == 'checked':
            methods_enabled.append('Rational Method (RM)')
        if hydro_vuln_responses.get('graphPeakCheckbox') == 'checked':
            methods_enabled.append('Graphical Peak Discharge Method (GPDM)')
        
        if methods_enabled:
            general_config['Analysis Methods'] = ', '.join(methods_enabled)
        
        # Return periods
        if 'rpListSelect[]' in hydro_vuln_responses:
            return_periods = hydro_vuln_responses['rpListSelect[]']
            if isinstance(return_periods, list):
                general_config['Return Periods (years)'] = ', '.join(map(str, return_periods))
            else:
                general_config['Return Periods (years)'] = str(return_periods)
        
        # Add general configuration table
        if general_config:
            gen_table = doc.add_table(rows=1, cols=2)
            gen_table.style = 'Table Grid'
            gen_table.columns[0].width = Inches(2.5)
            gen_table.columns[1].width = Inches(4.0)
            
            gen_headers = ['Configuration Parameter', 'Setting']
            gen_header_cells = gen_table.rows[0].cells
            for i, header in enumerate(gen_headers):
                gen_header_cells[i].text = header
                for paragraph in gen_header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
            
            for param, value in general_config.items():
                row_cells = gen_table.add_row().cells
                row_cells[0].text = param
                row_cells[1].text = value
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                            
        # A.2.2 Regional Frequency Analysis Parameters
        if hydro_vuln_responses.get('QregionalFrequencyCheckbox') == 'checked':
            doc.add_heading('A.2.2 Regional Frequency Analysis Parameters', 3)
            
            rfa_param_details = {
                'QtyofDistSelect': ('Streamflow Distribution Type', 'Probability distribution for streamflow frequency analysis'),
                'QhomTestSelect': ('Streamflow Homogeneity Test', 'Statistical test for regional homogeneity assessment'),
                'QoutlierMethodSelect': ('Streamflow Outlier Detection', 'Method for identifying and handling extreme values'),
                'QparEstSelect': ('Streamflow Parameter Estimation', 'Method for distribution parameter calculation'),
                'QNonStationarySelect': ('Streamflow Stationarity', 'Treatment of temporal trends in streamflow data'),
                'PItyofDistSelect': ('Precipitation Distribution Type', 'Probability distribution for precipitation frequency analysis'),
                'PIhomTestSelect': ('Precipitation Homogeneity Test', 'Statistical test for precipitation data consistency'),
                'PIoutlierMethodSelect': ('Precipitation Outlier Detection', 'Method for precipitation extreme value handling'),
                'PIparEstSelect': ('Precipitation Parameter Estimation', 'Method for precipitation distribution parameters'),
                'PINonStationarySelect': ('Precipitation Stationarity', 'Treatment of temporal trends in precipitation data')
            }
            
            # Only include Z-score parameters if outlier method is Z-score
            if hydro_vuln_responses.get('QoutlierMethodSelect') == 'Zscore':
                rfa_param_details['QzScoreSelect'] = ('Streamflow Z-Score Threshold', 'Statistical threshold for outlier identification')
            
            if hydro_vuln_responses.get('PIoutlierMethodSelect') == 'Zscore':
                rfa_param_details['PIzScoreSelect'] = ('Precipitation Z-Score Threshold', 'Statistical threshold for precipitation outliers')
            
            # Create RFA parameters table
            rfa_params = {k: v for k, v in hydro_vuln_responses.items() if k in rfa_param_details}
            
            if rfa_params:
                rfa_table = doc.add_table(rows=1, cols=3)
                rfa_table.style = 'Table Grid'
                rfa_table.columns[0].width = Inches(2.5)
                rfa_table.columns[1].width = Inches(1.0)
                rfa_table.columns[2].width = Inches(3.0)
                
                rfa_headers = ['Parameter', 'Selection', 'Description']
                rfa_header_cells = rfa_table.rows[0].cells
                for i, header in enumerate(rfa_headers):
                    rfa_header_cells[i].text = header
                    for paragraph in rfa_header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(11)
                
                for key, value in rfa_params.items():
                    if key in rfa_param_details:
                        row_cells = rfa_table.add_row().cells
                        param_name, description = rfa_param_details[key]
                        
                        row_cells[0].text = param_name
                        row_cells[1].text = format_statistical_parameter_value(key, value)
                        row_cells[2].text = description
                        
                        for cell in row_cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)
            
            # A.2.2.1 Data Sources by Gauged Watershed for RFA
            doc.add_heading('A.2.2.1 Data Sources by Gauged Watershed for RFA', 4)
            
            rfa_data_source_intro = """
    Regional Frequency Analysis requires watershed-specific streamflow and precipitation data. The following table documents the specific data sources used for each analyzed watershed.
    """
            doc.add_paragraph(rfa_data_source_intro.strip())
            
            # Extract watershed-specific data sources
            watersheds = hydro_vuln_responses.get('Gst_Names', [])
            if isinstance(watersheds, list) and len(watersheds) > 0:
                # Create data sources table with 5 columns
                rfa_data_table = doc.add_table(rows=1, cols=5)
                rfa_data_table.style = 'Table Grid'
                rfa_data_table.columns[0].width = Inches(1.5)  # Gauged Watershed
                rfa_data_table.columns[1].width = Inches(1.8)  # Streamflow Data
                rfa_data_table.columns[2].width = Inches(1.2)  # Streamflow Data Type
                rfa_data_table.columns[3].width = Inches(1.8)  # Precipitation Data
                rfa_data_table.columns[4].width = Inches(1.2)  # Precipitation Data Type
                
                rfa_data_headers = ['Gauged Watershed', 'Streamflow Data', 'Streamflow Data Type', 'Precipitation Data', 'Precipitation Data Type']
                rfa_data_header_cells = rfa_data_table.rows[0].cells
                for i, header in enumerate(rfa_data_headers):
                    rfa_data_header_cells[i].text = header
                    for paragraph in rfa_data_header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(11)
                
                # Add watershed-specific data information
                for ws in watersheds:
                    row_cells = rfa_data_table.add_row().cells
                    row_cells[0].text = ws
                    
                    # Streamflow data
                    streamflow_select_key = f'StreamflowDataSelect{ws}'
                    if streamflow_select_key in hydro_vuln_responses:
                        streamflow_type = hydro_vuln_responses[streamflow_select_key]
                        
                        # Display streamflow data source with proper formatting
                        if streamflow_type == 'ams_stream_series':
                            row_cells[1].text = 'Annual Maximum Series'
                            row_cells[2].text = 'Annual Maxima'
                        elif streamflow_type == 'full_stream_series':
                            row_cells[1].text = 'Full Stream Series'
                            row_cells[2].text = 'Full Series'
                        elif streamflow_type == 'streamflow_data_NA':
                            row_cells[1].text = 'No Data Available'
                            row_cells[2].text = 'N/A'
                        else:
                            row_cells[1].text = streamflow_type.replace('_', ' ').title()
                            row_cells[2].text = 'User Provided'
                    else:
                        row_cells[1].text = 'Not specified'
                        row_cells[2].text = 'N/A'
                    
                    # Precipitation data
                    precip_select_key = f'PrecipDataSelect{ws}'
                    if precip_select_key in hydro_vuln_responses:
                        precip_type = hydro_vuln_responses[precip_select_key]
                        
                        # Display precipitation data source with proper formatting
                        if precip_type == 'ams_precip_series':
                            row_cells[3].text = 'Annual Maximum Precipitation Series'
                            row_cells[4].text = 'Annual Maxima'
                        elif precip_type == 'full_precip_series':
                            row_cells[3].text = 'Full Precipitation Series'
                            row_cells[4].text = 'Full Series'
                        elif precip_type == 'precip_data_NA':
                            row_cells[3].text = 'No Data Available'
                            row_cells[4].text = 'N/A'
                        else:
                            row_cells[3].text = precip_type.replace('_', ' ').title()
                            row_cells[4].text = 'User Provided'
                    else:
                        row_cells[3].text = 'Not specified'
                        row_cells[4].text = 'N/A'
                    
                    # Format all cells
                    for cell in row_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)

        # A.2.2 Rational Method Parameters
        if hydro_vuln_responses.get('rationalMethodCheckbox') == 'checked':
            doc.add_heading('A.2.2 Rational Method Parameters', 3)
            
            rm_param_details = {
                'PIdataSelect': ('Precipitation Data Source', 'Source of precipitation intensity information'),
                'CoeffRunoffSelect': ('Runoff Coefficient Method', 'Approach for runoff coefficient determination'),
                'CulvertCapacitySelect': ('Culvert Capacity Method', 'Hydraulic calculation approach for infrastructure capacity')
            }
            
            # If using watershed specific rain gauges, include PI statistical parameters
            if hydro_vuln_responses.get('PIdataSelect') == 'From Watershed Specific Rain Gauges':
                pi_params = {
                    'PItyofDistSelect': ('Precipitation Distribution Type', 'Probability distribution for precipitation frequency analysis'),
                    'PIhomTestSelect': ('Precipitation Homogeneity Test', 'Statistical test for precipitation data consistency'),
                    'PIoutlierMethodSelect': ('Precipitation Outlier Detection', 'Method for precipitation extreme value handling'),
                    'PIparEstSelect': ('Precipitation Parameter Estimation', 'Method for precipitation distribution parameters'),
                    'PINonStationarySelect': ('Precipitation Stationarity', 'Treatment of temporal trends in precipitation data')
                }
                rm_param_details.update(pi_params)
                
                # Only include PI Z-score if outlier method is Z-score
                if hydro_vuln_responses.get('PIoutlierMethodSelect') == 'Zscore':
                    rm_param_details['PIzScoreSelect'] = ('Precipitation Z-Score Threshold', 'Statistical threshold for precipitation outliers')
            
            # Create RM parameters table
            rm_params = {k: v for k, v in hydro_vuln_responses.items() if k in rm_param_details}
            
            if rm_params:
                rm_table = doc.add_table(rows=1, cols=3)
                rm_table.style = 'Table Grid'
                rm_table.columns[0].width = Inches(2.5)
                rm_table.columns[1].width = Inches(1.0)
                rm_table.columns[2].width = Inches(3.0)
                
                rm_headers = ['Parameter', 'Selection', 'Description']
                rm_header_cells = rm_table.rows[0].cells
                for i, header in enumerate(rm_headers):
                    rm_header_cells[i].text = header
                    for paragraph in rm_header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(11)
                
                for key, value in rm_params.items():
                    if key in rm_param_details:
                        row_cells = rm_table.add_row().cells
                        param_name, description = rm_param_details[key]
                        
                        row_cells[0].text = param_name
                        row_cells[1].text = format_statistical_parameter_value(key, value)
                        row_cells[2].text = description
                        
                        for cell in row_cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)
    
    # ========================================
    # A.3 HYDROGEOMORPHOLOGIC VULNERABILITY PARAMETERS
    # ========================================
    if hydro_geo_responses:
        doc.add_heading('A.3 Hydrogeomorphologic Vulnerability Assessment Parameters', 2)
        
        hydrogeo_intro = """
Hydrogeomorphologic vulnerability parameters control the multi-criteria analysis procedures used to assess erosion processes, debris flow potential, and infrastructure susceptibility to geomorphologic hazards. Variable weights determine the relative importance of different environmental factors in composite vulnerability calculations.
"""
        doc.add_paragraph(hydrogeo_intro.strip())
        
        # A.3.1 Method Configuration
        doc.add_heading('A.3.1 Analysis Method Configuration', 3)
        
        method_config = {}
        if hydro_geo_responses.get('sbevaEnabled') == 'true':
            method_config['SBEVA (Streambank Erosion Vulnerability Assessment)'] = 'Enabled'
        if hydro_geo_responses.get('rusleEnabled') == 'true':
            method_config['RUSLE (Revised Universal Soil Loss Equation)'] = 'Enabled'  
        if hydro_geo_responses.get('weppEnabled') == 'true':
            method_config['WEPP (Water Erosion Prediction Project)'] = 'Enabled'
        if hydro_geo_responses.get('wdfmEnabled') == 'true':
            method_config['WDFM (Watershed Debris Flow Model)'] = 'Enabled'
        
        if method_config:
            method_table = doc.add_table(rows=1, cols=2)
            method_table.style = 'Table Grid'
            method_table.columns[0].width = Inches(4.0)
            method_table.columns[1].width = Inches(1.5)
            
            method_headers = ['Analysis Method', 'Status']
            method_header_cells = method_table.rows[0].cells
            for i, header in enumerate(method_headers):
                method_header_cells[i].text = header
                for paragraph in method_header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
            
            for method, status in method_config.items():
                row_cells = method_table.add_row().cells
                row_cells[0].text = method
                row_cells[1].text = status
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
        
        # A.3.2 SBEVA Parameters
        if hydro_geo_responses.get('sbevaEnabled') == 'true':
            doc.add_heading('A.3.2 SBEVA (Streambank Erosion Vulnerability Assessment) Parameters', 3)
            
            sbeva_param_details = {
                'SbevaStreamBufferDist': ('Stream Buffer Distance (m)', 'Analysis corridor width along stream channels'),
                'sbeva24hr100yrPIwt': ('24-hr 100-yr Precipitation Weight', 'Relative importance of extreme precipitation events'),
                'sbeva30yrPrecipNormwt': ('30-yr Precipitation Normal Weight', 'Relative importance of long-term precipitation patterns'),
                'sbeva30yrTempNormwt': ('30-yr Temperature Normal Weight', 'Relative importance of temperature regime'),
                'sbeva30yrSolarNormwt': ('30-yr Solar Radiation Weight', 'Relative importance of solar energy inputs'),
                'sbevaAWSrtznwt': ('Available Water Storage Weight', 'Relative importance of soil water holding capacity'),
                'sbevaRunoffClasswt': ('Runoff Classification Weight', 'Relative importance of surface runoff characteristics'),
                'sbevaDrainageClasswt': ('Drainage Classification Weight', 'Relative importance of soil drainage properties'),
                'sbevaHydroSoilGrpwt': ('Hydrologic Soil Group Weight', 'Relative importance of soil infiltration capacity'),
                'sbevaSlopewt': ('Slope Weight', 'Relative importance of topographic gradient'),
                'sbevaLandCoverwt': ('Land Cover Weight', 'Relative importance of vegetation and surface cover')
            }
            
            sbeva_params = {k: v for k, v in hydro_geo_responses.items() if k in sbeva_param_details}
            
            if sbeva_params:
                sbeva_table = doc.add_table(rows=1, cols=3)
                sbeva_table.style = 'Table Grid'
                sbeva_table.columns[0].width = Inches(2.5)
                sbeva_table.columns[1].width = Inches(1.0)
                sbeva_table.columns[2].width = Inches(3.0)
                
                sbeva_headers = ['Variable', 'Weight', 'Description']
                sbeva_header_cells = sbeva_table.rows[0].cells
                for i, header in enumerate(sbeva_headers):
                    sbeva_header_cells[i].text = header
                    for paragraph in sbeva_header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(11)
                
                for key, value in sbeva_params.items():
                    if key in sbeva_param_details:
                        row_cells = sbeva_table.add_row().cells
                        param_name, description = sbeva_param_details[key]
                        
                        row_cells[0].text = param_name
                        row_cells[1].text = format_parameter_value(value)
                        row_cells[2].text = description
                        
                        for cell in row_cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)
        
        # A.3.3 WDFM Parameters
        if hydro_geo_responses.get('wdfmEnabled') == 'true':
            doc.add_heading('A.3.3 WDFM (Watershed Debris Flow Model) Parameters', 3)
            
            wdfm_param_details = {
                'wdfmAwswt': ('Available Water Storage Weight', 'Soil water retention influence on debris flow initiation'),
                'wdfmDrainageClasswt': ('Drainage Classification Weight', 'Soil drainage impact on slope stability'),
                'wdfmKfactwt': ('K-Factor Weight', 'Soil erodibility contribution to debris mobilization'),
                'wdfmKsatwt': ('Saturated Hydraulic Conductivity Weight', 'Soil permeability effect on pore pressure'),
                'wdfmRunoffClasswt': ('Runoff Classification Weight', 'Surface flow contribution to debris triggering'),
                'wdfmSoilSlipwt': ('Soil Slip Potential Weight', 'Inherent slope instability characteristics'),
                'wdfmSoilBtmDepthwt': ('Soil Bottom Depth Weight', 'Soil profile thickness influence on stability'),
                'wdfmSoilTaxonomicwt': ('Soil Taxonomic Weight', 'Soil classification impact on debris flow susceptibility'),
                'wdfmTFactorwt': ('T-Factor Weight', 'Soil loss tolerance effect on long-term stability'),
                'wdfmGeologyRoackType1wt': ('Geology Rock Type Weight', 'Bedrock geology influence on slope stability'),
                'wdfmNDVIwt': ('NDVI Weight', 'Vegetation cover effect on slope reinforcement'),
                'wdfmPIwt': ('Precipitation Intensity Weight', 'Rainfall trigger intensity for debris flows'),
                'wdfmSlopewt': ('Slope Weight', 'Topographic gradient contribution to debris flow potential'),
                'wdfmRoadBufferwt': ('Road Buffer Weight', 'Transportation infrastructure proximity effects'),
                'wdfmStreamBufferwt': ('Stream Buffer Weight', 'Channel network proximity influence')
            }
            
            wdfm_params = {k: v for k, v in hydro_geo_responses.items() if k in wdfm_param_details}
            
            if wdfm_params:
                wdfm_table = doc.add_table(rows=1, cols=3)
                wdfm_table.style = 'Table Grid'
                wdfm_table.columns[0].width = Inches(2.5)
                wdfm_table.columns[1].width = Inches(1.0)
                wdfm_table.columns[2].width = Inches(3.0)
                
                wdfm_headers = ['Variable', 'Weight', 'Description']
                wdfm_header_cells = wdfm_table.rows[0].cells
                for i, header in enumerate(wdfm_headers):
                    wdfm_header_cells[i].text = header
                    for paragraph in wdfm_header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(11)
                
                for key, value in wdfm_params.items():
                    if key in wdfm_param_details:
                        row_cells = wdfm_table.add_row().cells
                        param_name, description = wdfm_param_details[key]
                        
                        row_cells[0].text = param_name
                        row_cells[1].text = format_parameter_value(value)
                        row_cells[2].text = description
                        
                        for cell in row_cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)
    
    # ========================================
    # APPENDIX B: DESCRIPTION OF METHODOLOGIES
    # ========================================
    doc.add_heading('Appendix B: Description of Methodologies', 1)

    method_intro = """
    This appendix provides detailed descriptions of the analytical methodologies employed in the CULVERT assessment. Only methods that were enabled and executed in this specific analysis are documented below.
    """
    doc.add_paragraph(method_intro.strip())

    # B.1 Watershed Delineation Methodology
    if ws_deln_responses:
        doc.add_heading('B.1 Watershed Delineation Methodology', 2)
        
        ws_method_text = """
    Watershed delineation forms the foundation of hydrologic analysis by defining the spatial boundaries within which surface water flows converge toward a common outlet point. The methodology employs advanced geospatial processing techniques to accurately determine drainage areas, flow paths, and watershed characteristics essential for flood modeling and infrastructure risk assessment.

    The process begins with preprocessing of user-provided digital elevation models (DEM) through hydrological conditioning procedures. Depression filling and sink removal operations ensure continuous flow routing across the terrain surface. Flow direction calculations utilize the deterministic eight-direction (D8) algorithm to establish the steepest descent path from each grid cell to its neighbors.

    Flow accumulation analysis determines the upstream contributing area for each cell by tracing flow paths from ridge lines to valley bottoms. Stream network delineation applies user-defined flow accumulation thresholds to identify channel initiation points and establish the drainage network topology. Pour points representing road-stream crossing locations are precisely positioned through automated snapping procedures that align infrastructure points with the computed stream network.

    When hydro-enforcement is required, the methodology implements advanced terrain modification techniques to address limitations in elevation data representation of drainage systems beneath transportation infrastructure. Road surfaces are elevated above surrounding terrain to prevent artificial flow barriers, while strategic breaklines are burned into the DEM along known drainage paths to ensure proper flow connectivity.

    The final watershed delineation process traces upslope contributing areas from each pour point using recursive flow path analysis. Critical watershed parameters including drainage area, average slope, main channel length, and time of concentration are computed from the delineated boundaries and integrated flow network. These parameters provide essential inputs for subsequent hydrologic vulnerability assessments and peak discharge calculations.
    """
        doc.add_paragraph(ws_method_text.strip())

    # B.2 Hydrologic Vulnerability Assessment Methodologies
    if hydro_vuln_responses:
        doc.add_heading('B.2 Hydrologic Vulnerability Assessment Methodologies', 2)
        
        # Check which methods were enabled
        if hydro_vuln_responses.get('QregionalFrequencyCheckbox') == 'checked':
            doc.add_heading('B.2.1 Regional Frequency Analysis (RFA)', 3)
            
            rfa_text = """
    Regional Frequency Analysis addresses the fundamental challenge of estimating flood frequencies at ungauged locations or sites with limited streamflow records. The methodology combines statistical data from multiple gauging stations within homogeneous regions to develop robust frequency relationships that account for both spatial variability and temporal trends in extreme hydrologic events.

    The analysis begins with comprehensive data compilation from available streamflow records within the study region. Annual maximum series are extracted from continuous discharge records and subjected to rigorous quality control procedures including outlier detection, trend analysis, and homogeneity testing. The L-moments method provides robust parameter estimation for probability distributions, offering superior performance compared to conventional moment-based approaches, particularly for extreme value analysis.

    Regional homogeneity assessment employs the heterogeneity measure H-statistic to evaluate the appropriateness of combining data from different sites. Discordancy measures identify stations with unusual flood characteristics that may require separate treatment or exclusion from regional analysis. The methodology accommodates both stationary and non-stationary conditions through trend-aware frequency analysis when significant temporal patterns are detected.

    Distribution selection follows a systematic approach comparing multiple candidate distributions including Generalized Extreme Value (GEV), Pearson Type III, and Log-Pearson Type III. Model performance is evaluated using L-moment ratio diagrams, goodness-of-fit tests, and regional growth factor analysis. Bootstrap resampling techniques provide confidence intervals for frequency estimates, enabling uncertainty quantification in design applications.

    The final regional frequency relationships express flood quantiles as functions of watershed characteristics and return periods. These relationships enable flood frequency estimation at ungauged locations within the homogeneous region, supporting infrastructure design and risk assessment applications with quantified uncertainty bounds.
    """
            doc.add_paragraph(rfa_text.strip())
        
        if hydro_vuln_responses.get('rationalMethodCheckbox') == 'checked':
            doc.add_heading('B.2.2 Rational Method', 3)
            
            rm_text = """
    The Rational Method provides a computationally efficient approach for peak discharge estimation based on the fundamental relationship between rainfall intensity, watershed characteristics, and runoff generation processes. The methodology applies the classical equation Q = C × I × A, where Q represents peak discharge, C is the runoff coefficient, I is rainfall intensity, and A is the drainage area.

    Runoff coefficient determination integrates multiple environmental factors that influence the proportion of precipitation converted to surface runoff. Land use classification utilizes the National Land Cover Database (NLCD) to identify surface cover types ranging from impervious urban areas to forested landscapes. Hydrologic soil group classification from the gridded Soil Survey Geographic (gSSURGO) database categorizes soils based on infiltration characteristics and drainage properties.

    Topographic analysis extracts average watershed slope from digital elevation models to account for terrain effects on runoff velocity and concentration time. The methodology accommodates varying data availability scenarios through hierarchical coefficient determination approaches. When precipitation-streamflow relationships are available, back-calculation procedures derive site-specific coefficients that reflect local hydrologic conditions.

    Rainfall intensity calculations utilize precipitation frequency data from NOAA Atlas-14 or user-provided intensity-duration-frequency relationships. The methodology applies appropriate duration selection based on watershed time of concentration, ensuring consistency between rainfall timing and basin response characteristics. Correction factors account for spatial rainfall distribution and temporal variability effects on peak discharge generation.

    Time of concentration estimation employs the SCS methodology, combining overland flow travel time with channel flow routing calculations. The methodology considers land cover effects on surface roughness, slope influences on flow velocity, and channel geometry impacts on hydraulic routing. Final peak discharge calculations integrate all components through the rational equation framework, providing rapid assessment capabilities suitable for preliminary design and screening-level analysis applications.
    """
            doc.add_paragraph(rm_text.strip())
        
        if hydro_vuln_responses.get('graphPeakCheckbox') == 'checked':
            doc.add_heading('B.2.3 Graphical Peak Discharge Method (GPDM)', 3)
            
            gpdm_text = """
    The Graphical Peak Discharge Method implements the standardized TR-55 approach developed by the USDA Natural Resources Conservation Service for estimating peak discharge from small watersheds. The methodology combines the SCS Curve Number technique with unit hydrograph theory to provide reliable peak flow estimates for ungauged watersheds with drainage areas typically less than 25 square kilometers.

    Curve Number determination integrates land use classification from the National Land Cover Database with hydrologic soil group data to characterize the runoff potential of different watershed areas. The methodology accounts for antecedent moisture conditions through appropriate curve number adjustments that reflect soil saturation effects on runoff generation. Composite curve numbers are calculated as area-weighted averages across heterogeneous watersheds.

    Initial abstraction calculations utilize the standard relationship Ia = 0.2S, where S represents the maximum potential retention calculated from curve numbers. The methodology accommodates alternative initial abstraction ratios when local calibration data supports different relationships. Precipitation depth selection utilizes NOAA Atlas-14 point precipitation frequency estimates with appropriate areal reduction factors for larger watersheds.

    Time of concentration calculations follow SCS procedures for combining overland flow and channel flow components. The methodology considers land cover effects on surface roughness, channel geometry influences on hydraulic routing, and pond/wetland impacts on flow attenuation. Rainfall distribution selection utilizes regional Type I, IA, II, or III distributions based on geographic location and climatic characteristics.

    Unit peak discharge determination employs TR-55 graphical relationships that express peak flow as a function of curve number, time of concentration, and initial abstraction ratio. The methodology interpolates between tabulated values to provide continuous relationships suitable for diverse watershed conditions. Final peak discharge calculations multiply unit peak discharge by drainage area and precipitation depth, incorporating appropriate adjustment factors for pond and wetland storage effects.
    """
            doc.add_paragraph(gpdm_text.strip())

    # B.3 Hydrogeomorphologic Vulnerability Assessment Methodologies
    if hydro_geo_responses:
        doc.add_heading('B.3 Hydrogeomorphologic Vulnerability Assessment Methodologies', 2)
        
        if hydro_geo_responses.get('sbevaEnabled') == 'true':
            doc.add_heading('B.3.1 Streambank Erosion Vulnerability Assessment (SBEVA)', 3)
            
            sbeva_text = """
    The Streambank Erosion Vulnerability Assessment (SBEVA) methodology evaluates the susceptibility of stream channels to lateral erosion processes that can threaten adjacent transportation infrastructure through bank undermining, channel widening, and flow pattern alteration. The approach integrates multiple environmental factors through a multi-criteria analysis framework that synthesizes climatic, topographic, soil, and vegetation controls on erosion processes.

    Climatic vulnerability factors incorporate precipitation intensity data from NOAA Atlas-14 to quantify erosive forcing from extreme rainfall events. Temperature and solar radiation data from PRISM climate normals characterize the energy environment that influences vegetation growth, soil moisture dynamics, and freeze-thaw processes affecting bank stability. The methodology weights extreme precipitation events more heavily than average conditions to emphasize infrastructure threats from high-magnitude erosive forces.

    Soil property analysis utilizes the gridded Soil Survey Geographic (gSSURGO) database to extract parameters controlling erosion resistance and hydrologic response. Available water storage capacity indicates soil moisture retention characteristics that influence bank saturation and stability. Drainage class and hydrologic soil group classifications characterize subsurface flow patterns and infiltration rates that affect pore pressure development and bank failure mechanisms.

    Topographic factors integrate slope gradients calculated from digital elevation models to quantify gravitational forces promoting bank instability. Stream channel proximity analysis establishes spatial buffers around the drainage network to focus assessment on areas directly affecting transportation infrastructure. The methodology considers both immediate channel banks and adjacent slopes that contribute sediment and debris to the stream system.

    Vegetation analysis utilizes land cover classification from the National Land Cover Database to evaluate protective effects from root reinforcement, surface armoring, and flow resistance. Forest cover provides maximum protection through deep root systems and surface roughness, while agricultural and developed areas offer minimal erosion resistance. The methodology applies scientifically-based weighting factors that reflect empirical relationships between vegetation characteristics and erosion rates.

    Multi-criteria integration employs weighted overlay procedures where user-defined variable weights reflect relative importance of different environmental factors. Standardization procedures ensure compatibility among variables with different units and scales. The final SBEVA vulnerability classification ranges from Very Low to Very High, providing actionable information for infrastructure management and maintenance prioritization decisions.
    """
            doc.add_paragraph(sbeva_text.strip())
        
        if hydro_geo_responses.get('rusleEnabled') == 'true':
            doc.add_heading('B.3.2 Revised Universal Soil Loss Equation (RUSLE)', 3)
            
            rusle_text = """
    The Revised Universal Soil Loss Equation (RUSLE) methodology quantifies long-term average annual soil erosion rates across watershed surfaces through integration of climatic, topographic, soil, vegetation, and management factors. The approach provides spatially distributed erosion estimates that identify sediment source areas and quantify delivery potential to stream channels where sediment accumulation may impact infrastructure performance.

    Rainfall erosivity factor (R) calculations utilize precipitation intensity data to quantify the kinetic energy available for soil detachment and transport. The methodology applies established relationships between rainfall intensity and erosive power, accounting for storm duration and frequency characteristics. Regional R-factor maps provide standardized erosivity values calibrated to local climatic conditions, ensuring appropriate representation of erosive forcing across diverse geographic regions.

    Soil erodibility factor (K) determination utilizes soil property data from the gridded Soil Survey Geographic database to characterize intrinsic resistance to erosion. The methodology considers particle size distribution, organic matter content, soil structure, and permeability factors that control susceptibility to rainfall and runoff erosion. K-factor values reflect extensive field research quantifying erosion rates under standardized plot conditions.

    Topographic factor (LS) calculations integrate slope length and steepness effects on erosion processes through terrain analysis of digital elevation models. Flow accumulation algorithms determine upslope contributing areas that influence runoff erosivity, while slope gradient calculations quantify gravitational forces promoting soil detachment. The methodology applies established LS-factor equations that account for both rill and interrill erosion processes.

    Cover and management factor (C) estimation utilizes vegetation indices derived from satellite imagery to characterize protective effects from plant cover. The methodology applies relationships between Normalized Difference Vegetation Index (NDVI) values and ground cover percentages to estimate C-factors across diverse land use types. Seasonal variations in vegetation cover are accommodated through multi-temporal imagery analysis when available.

    Support practice factor (P) accounts for conservation measures such as contouring, strip cropping, and terracing that reduce erosion through altered flow patterns and sediment trapping. The methodology applies standard P-factor values from RUSLE documentation, modified based on local topographic and management conditions when detailed information is available.

    Final erosion rate calculations apply the classical RUSLE structure: A = R × K × LS × C × P, where A represents average annual soil loss. Spatial analysis produces distributed erosion maps that identify critical source areas and quantify sediment delivery potential to transportation infrastructure. Results support prioritization of erosion control measures and assessment of long-term maintenance requirements for road-stream crossing systems.
    """
            doc.add_paragraph(rusle_text.strip())
        
        if hydro_geo_responses.get('wdfmEnabled') == 'true':
            doc.add_heading('B.3.3 Watershed Debris Flow Model (WDFM)', 3)
            
            wdfm_text = """
    The Watershed Debris Flow Model (WDFM) methodology assesses mass wasting hazards that pose significant threats to transportation infrastructure through debris delivery, impact forces, and channel blockage during extreme precipitation events. The approach integrates slope stability analysis with debris flow initiation and runout modeling to identify vulnerable infrastructure locations and quantify relative hazard levels across diverse terrain conditions.

    Slope stability assessment utilizes digital elevation models to calculate slope gradients and identify terrain positions susceptible to shallow landsliding and debris flow initiation. The methodology considers threshold slope angles based on regional geologic and climatic conditions, with steeper slopes receiving higher vulnerability scores. Convergent topography and channel head locations receive additional weighting due to flow concentration effects that promote debris mobilization.

    Soil property analysis incorporates available water storage, drainage characteristics, and saturated hydraulic conductivity from the gridded Soil Survey Geographic database to evaluate subsurface conditions affecting slope stability. High water storage capacity combined with poor drainage creates conditions conducive to pore pressure development and slope failure during intense precipitation events. Soil taxonomic classifications provide additional constraints on debris flow susceptibility based on established relationships between soil types and mass wasting processes.

    Geological factors integrate bedrock lithology and surficial deposit characteristics that control slope stability and debris supply potential. The methodology utilizes geological map data to identify rock types associated with high weathering rates, weak bedding planes, or structural discontinuities that promote slope instability. Volcanic terrains, weak sedimentary rocks, and highly fractured crystalline rocks typically receive elevated vulnerability ratings.

    Vegetation analysis evaluates protective effects from root reinforcement and surface stabilization using land cover classification and vegetation indices. Forest cover provides maximum slope protection through deep root systems that mechanically reinforce soil profiles. The methodology accounts for vegetation mortality from disturbances such as wildfire, insect outbreaks, or timber harvest that temporarily reduce slope stability and increase debris flow susceptibility.

    Precipitation trigger analysis incorporates extreme rainfall intensities from NOAA Atlas-14 to identify threshold conditions for debris flow initiation. The methodology considers both short-duration high-intensity events and longer-duration moderate-intensity storms that promote slope saturation. Regional calibration with historical debris flow occurrence data refines precipitation thresholds for local conditions when available.

    Infrastructure proximity analysis establishes spatial buffers around transportation networks to focus assessment on areas directly threatening road and bridge systems. The methodology considers both impact zones where debris flows may directly strike infrastructure and deposition zones where debris accumulation may block drainage systems or alter flow patterns. Channel network analysis identifies flow paths that connect potential source areas with infrastructure locations.

    Multi-criteria integration employs weighted overlay procedures similar to SBEVA methodology, with user-defined weights reflecting relative importance of different hazard factors. Vulnerability classifications range from Very Low to Very High based on composite scores that integrate all assessment factors. Results provide spatially explicit hazard maps supporting infrastructure risk management, emergency planning, and adaptive design decisions for transportation systems in mountainous terrain.
    """
            doc.add_paragraph(wdfm_text.strip())
    
    
    # Check if more than one method was enabled for EHVI
    enabled_methods = []
    if hydro_geo_responses.get('sbevaEnabled') == 'true':
        enabled_methods.append('SBEVA')
    if hydro_geo_responses.get('rusleEnabled') == 'true':
        enabled_methods.append('RUSLE')
    if hydro_geo_responses.get('weppEnabled') == 'true':
        enabled_methods.append('WEPP')
    if hydro_geo_responses.get('wdfmEnabled') == 'true':
        enabled_methods.append('WDFM')
    
    # Add EHVI section if more than one method was enabled
    if len(enabled_methods) > 1:
        doc.add_heading('B.3.4 Ensemble Hydrogeomorphologic Vulnerability Index (EHVI)', 3)
        
        methods_text = ', '.join(enabled_methods)
        ehvi_text = f"""
The Ensemble Hydrogeomorphologic Vulnerability Index (EHVI) provides comprehensive risk synthesis by integrating multiple erosion and mass wasting vulnerability assessments into a unified decision-support metric. This analysis combines results from {len(enabled_methods)} analytical methods ({methods_text}) to deliver a holistic vulnerability assessment that captures complex interactions between multiple geomorphologic processes threatening transportation infrastructure.

The ensemble methodology addresses limitations inherent in individual assessment approaches by leveraging the complementary strengths of different analytical frameworks. While individual methods may excel in specific process domains or environmental conditions, the ensemble approach provides robust vulnerability characterization across diverse terrain and climatic settings. This integrated perspective supports more informed infrastructure management decisions by reducing uncertainty and providing comprehensive risk evaluation.

Methodological integration employs standardized vulnerability scaling procedures that ensure compatibility among different analytical approaches. Each component method generates vulnerability classifications ranging from Very Low to Very High based on method-specific criteria and thresholds. The EHVI framework normalizes these classifications to common scales and applies equal weighting to each component method, reflecting the assumption that different geomorphologic processes contribute equally to overall infrastructure vulnerability.

Statistical aggregation procedures combine normalized vulnerability scores through averaging operations that preserve the full range of vulnerability levels while providing smoothed composite assessments. The methodology accommodates missing data scenarios where individual methods may not provide coverage for all watershed areas due to data limitations or methodological constraints. Spatial interpolation and gap-filling procedures ensure complete coverage across the analysis domain.

Uncertainty quantification within the EHVI framework considers both methodological uncertainty from individual component models and integration uncertainty from the ensemble combination process. Variance measures across component methods provide indicators of assessment confidence, with high inter-method agreement indicating robust vulnerability estimates and high variance suggesting areas requiring additional investigation or alternative analytical approaches.

The final EHVI classification maintains the five-category vulnerability scale (Very Low, Low, Moderate, High, Very High) established by component methods while providing enhanced discrimination through ensemble averaging effects. Results support prioritization of infrastructure improvements, maintenance scheduling, and adaptive management strategies by identifying locations where multiple geomorphologic processes converge to create elevated threat levels.

Validation and calibration of EHVI results utilize available field observations, historical damage records, and expert knowledge to assess ensemble performance relative to component methods. The methodology supports adaptive weighting schemes where local experience or additional data support differential importance among component processes, enabling customization for specific regional or project conditions while maintaining scientific rigor and methodological transparency.
"""
        doc.add_paragraph(ehvi_text.strip())
    # ========================================
    # APPENDIX C: REFERENCES AND CITATIONS
    # ========================================
    doc.add_heading('Appendix C: References and Citations', 1)
    
    # ========================================
    # APPENDIX D: GLOSSARY OF TERMS
    # ========================================
    doc.add_heading('Appendix D: Glossary of Technical Terms', 1)
    
    glossary_intro = """
This glossary defines technical terms, abbreviations, and specialized vocabulary used throughout the CULVERT analysis and report. Definitions are provided to ensure clear understanding of methodological concepts and analytical results for users with varying technical backgrounds.
"""
    doc.add_paragraph(glossary_intro.strip())
    
    # Create comprehensive glossary
    glossary_terms = {
        "Annual Maximum Series (AMS)": "Statistical dataset consisting of the largest flood peak or precipitation value recorded in each year of the observation period, used for frequency analysis of extreme events.",
        
        "Bootstrap Method": "Statistical resampling technique that estimates the sampling distribution of a statistic by repeatedly sampling with replacement from the original data, used for confidence interval estimation.",
        
        "Curve Number": "Dimensionless parameter used in the SCS Curve Number method to represent the runoff potential of different land use and soil combinations, ranging from 30 (low runoff) to 100 (high runoff).",
        
        "Debris Flow": "Fast-moving landslide consisting of water-saturated debris that flows down slopes and channels, posing significant hazards to infrastructure through impact forces and blockage.",
        
        "DEM (Digital Elevation Model)": "Raster dataset representing ground surface elevations at regular spatial intervals, fundamental input for terrain analysis and hydrological modeling.",
        
        "EHVI (Ensemble Hydrogeomorphological Vulnerability Index)": "Composite vulnerability metric integrating results from multiple erosion assessment methods (SBEVA, RUSLE, WDFM) to provide comprehensive geomorphologic risk evaluation.",
        
        "Flow Accumulation": "Raster dataset representing the accumulated upstream drainage area contributing to each cell, calculated from flow direction and used for stream network delineation.",
        
        "Frequency Analysis": "Statistical method for estimating the probability of occurrence and return periods of extreme hydrological events such as floods, droughts, or storms.",
        
        "GPDM (Graphical Peak Discharge Method)": "Engineering method for estimating peak discharge using graphical relationships between watershed characteristics, precipitation, and runoff, typically following TR-55 procedures.",
        
        "Heterogeneity": "Measure of regional variability in frequency analysis, quantifying how similar different sites are within a homogeneous region using statistical tests.",
        
        "Hydrograph": "Graph showing the variation of streamflow discharge over time, typically displaying the response of a watershed to precipitation events.",
        
        "L-moments": "Linear combinations of probability-weighted moments used for parameter estimation in frequency analysis, providing robust estimates for extreme value distributions.",
        
        "Non-stationarity": "Condition where statistical properties of a time series (mean, variance, distribution) change over time, often due to climate change or human activities.",
        
        "Pour Point": "Specific location in a watershed, typically representing a road-stream crossing, where water flows converge and discharge calculations are performed.",
        
        "Return Period": "Average time interval between occurrences of a hydrological event of given magnitude or greater, expressed in years (e.g., 100-year flood).",
        
        "RFA (Regional Frequency Analysis)": "Statistical method that combines data from multiple sites within a homogeneous region to improve frequency estimates, particularly for ungauged locations.",
        
        "RM (Rational Method)": "Empirical method for calculating peak discharge using the equation Q = CiA, where Q is discharge, C is runoff coefficient, i is rainfall intensity, and A is drainage area.",
        
        "RUSLE (Revised Universal Soil Loss Equation)": "Empirical model for predicting long-term average annual soil erosion rates based on rainfall erosivity, soil erodibility, topography, vegetation, and management practices.",
        
        "SBEVA (Streambank Erosion Vulnerability Assessment)": "Multi-criteria analysis method for evaluating the susceptibility of stream channels to lateral erosion processes affecting adjacent infrastructure.",
        
        "Time of Concentration": "Time required for runoff to travel from the most distant point in a watershed to the outlet, critical parameter for peak discharge calculations.",
        
        "TR-55": "Technical Release 55 published by USDA-NRCS, providing standardized methods for urban hydrology calculations including the Graphical Peak Discharge Method.",
        
        "UTM (Universal Transverse Mercator)": "Map projection system dividing the Earth into 60 zones, providing accurate distance and area calculations for regional-scale analysis.",
        
        "Vulnerability Classification": "Categorical system (typically Very Low, Low, Moderate, High, Very High) used to classify infrastructure risk levels based on quantitative vulnerability assessments.",
        
        "Watershed Delineation": "Process of determining the drainage boundary and area contributing surface runoff to a specific outlet point, fundamental step in hydrological analysis.",
        
        "WDFM (Watershed Debris Flow Model)": "Multi-criteria assessment method for evaluating debris flow susceptibility based on topographic, geological, soil, vegetation, and climatic factors."
    }
    
    # Add glossary terms in alphabetical order
    for term, definition in sorted(glossary_terms.items()):
        term_para = doc.add_paragraph()
        term_run = term_para.add_run(f"{term}: ")
        term_run.bold = True
        term_run.font.size = Pt(11)
        
        def_run = term_para.add_run(definition)
        def_run.font.size = Pt(10)
        
        # Add small spacing between terms
        term_para.space_after = Pt(6)
    
    # Add final note
    doc.add_paragraph()
    final_note = doc.add_paragraph()
    final_note_run = final_note.add_run("Note: ")
    final_note_run.bold = True
    final_note_run.font.size = Pt(10)
    
    final_note_text = final_note.add_run("This glossary provides definitions specific to the CULVERT Web Application context. For additional technical terms or detailed methodological explanations, readers should consult the primary references listed in Appendix C.")
    final_note_text.font.size = Pt(10)
    final_note_text.italic = True

def format_parameter_value(value):
    """Helper function to format parameter values for display"""
    if isinstance(value, list):
        return ', '.join(map(str, value))
    elif isinstance(value, (int, float)):
        if isinstance(value, float):
            return f"{value:.3f}"
        else:
            return str(value)
    elif isinstance(value, bool):
        return "Yes" if value else "No"
    elif value is None:
        return "Not specified"
    else:
        return str(value).replace('_', ' ').title()